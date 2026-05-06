"""
Append verified references section to T5 v4 manuscript.

Reads the verified citation list from T5_build_full_lit_verification.py CITED
constant, builds a numbered Vancouver-style References section grouped by
lens, and appends to T5_title_intro_methods_results_discussion_package_v4_with_figures.docx,
saving as v4.1_with_references.docx.

This addresses the gap that v4_with_figures has 14+ author-narrative mentions
(Taylor, EMP, Shoemaker, Hubbell, Fisher, Preston, Grilli, Cao, etc.) but no
References section.
"""
from __future__ import annotations
from copy import deepcopy
from pathlib import Path
import importlib.util
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/ynh83/Desktop/T5_Macroecology")
SRC = ROOT / "T5_title_intro_methods_results_discussion_package_v4_with_figures.docx"
OUT = ROOT / "T5_title_intro_methods_results_discussion_package_v4.1_with_references.docx"

LIT_SCRIPT = ROOT / "scripts" / "T5_build_full_lit_verification.py"
spec = importlib.util.spec_from_file_location("lit_verif", LIT_SCRIPT)
lit_mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(lit_mod)
CITED = lit_mod.CITED

doc = Document(str(SRC))

# Add a clear page break before References for clean print layout
doc.add_page_break()

heading = doc.add_heading("References", level=1)
for run in heading.runs:
    run.font.name = "Times New Roman"

note = doc.add_paragraph()
r = note.add_run(
    "All identifiers verified via PubMed esummary and CrossRef on 2026-05-04. "
    "Full audit table with clickable verification links is available at "
    "T5_References_FullField_Verification_2026-05-04.docx."
)
r.italic = True
r.font.size = Pt(10.5)
r.font.name = "Times New Roman"

current_lens = None
for entry in CITED:
    idx, lens, citation, identifier, status, link = entry
    if lens != current_lens:
        # add a sub-heading per lens for grouped readability
        sub = doc.add_paragraph()
        sr = sub.add_run(lens)
        sr.bold = True
        sr.font.size = Pt(11)
        sr.font.name = "Times New Roman"
        current_lens = lens

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.first_line_indent = Pt(-18)
    para.paragraph_format.space_after = Pt(2)

    # Reference number
    num_run = para.add_run(f"{idx}. ")
    num_run.bold = True
    num_run.font.size = Pt(10.5)
    num_run.font.name = "Times New Roman"

    # Citation text
    cit_run = para.add_run(citation)
    cit_run.font.size = Pt(10.5)
    cit_run.font.name = "Times New Roman"

    # Identifier + status (small)
    meta_run = para.add_run(f"  [{identifier}; {status}]")
    meta_run.italic = True
    meta_run.font.size = Pt(9.5)
    meta_run.font.name = "Times New Roman"

doc.save(str(OUT))
print(f"Wrote: {OUT}")
print(f"  size: {OUT.stat().st_size:,} bytes")
print(f"  references appended: {len(CITED)}")
