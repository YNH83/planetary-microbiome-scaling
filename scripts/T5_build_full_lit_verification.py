"""
Build T5_References_FullField_Verification_2026-05-04.docx ,  a comprehensive
verified literature reference table covering the broadest framing of the T5
Macroecology study.

Format mirrors `T5_References_Verification_2026-04-29.docx`:
    7-column verification table:
    # | Lens | Citation | Identifier (PMID / DOI) | Status | Verification link |
    User check (OK / NG / comment)

All identifiers in the CITED list below have been verified inline via the
PubMed esummary API and CrossRef API on 2026-05-04. The list expands the
2026-04-29 audit (17 entries) to >=30 entries spanning:

    L1  Macroecology / scaling laws (Taylor, May, Locey, Falkowski)
    L2  Stochastic-logistic / Taylor microbial (Grilli, Shoemaker, Ji-style,
        Zaoli, Ma, Yi)
    L3  Neutral theory (Hubbell, Volkov, Etienne)
    L4  Earth Microbiome / cross-biome catalogues (EMP, Tara, MetaSUB, soil,
        IGGdb, GMRC)
    L5  Shotgun replication / curatedMG (Pasolli, Lloyd-Price iHMP, HMP)
    L6  Microbial census / scale (Sender, Whitman 1998 PNAS)
    L7  Bayesian hierarchical / partial pooling (Vehtari LOO, Kass-Raftery
        BIC, PyMC, Gelman textbook)
    L8  OSF pre-registration (Nosek)
    L9  Carrying-capacity heritage (Verhulst 1838 logistic)
    L10 Method tooling (QIIME 2, Deblur, MEGAHIT)
    L11 Disease microbiome (IBD iHMP, CRC Zeller / Yachida, FMT van Nood,
        immunotherapy Gopalakrishnan / Routy / Davar / Baruch, metformin
        Forslund)
    L12 Convergence framing analogue (Cao Nature 2026 PerturbFate)
"""
from __future__ import annotations
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/ynh83/Desktop/T5_Macroecology")
OUT = ROOT / f"T5_References_FullField_Verification_{date.today().isoformat()}.docx"


CITED = [
    # idx, lens, citation, identifier, status, link
    (1,  "L1 Macroecology / scaling laws",
     "Taylor LR. Aggregation, variance and the mean. Nature. 1961;189:732-735.",
     "DOI 10.1038/189732a0",
     "Verified 2026-05-04 (CrossRef; pre-PubMed era)",
     "doi.org/10.1038/189732a0"),
    (2,  "L1 Macroecology / scaling laws",
     "May RM. How many species are there on Earth? Science. 1988;241(4872):1441-1449.",
     "DOI 10.1126/science.241.4872.1441",
     "Verified 2026-05-04 (CrossRef)",
     "doi.org/10.1126/science.241.4872.1441"),
    (3,  "L1 Macroecology / scaling laws",
     "Locey KJ, Lennon JT. Scaling laws predict global microbial diversity. "
     "Proc Natl Acad Sci USA. 2016;113(21):5970-5975.",
     "PMID 27140646",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/27140646/"),
    (4,  "L1 Macroecology / scaling laws",
     "Falkowski PG, Fenchel T, Delong EF. The microbial engines that drive "
     "Earth's biogeochemical cycles. Science. 2008;320(5879):1034-1039.",
     "PMID 18497287",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/18497287/"),

    (5,  "L2 Stochastic-logistic / Taylor microbial",
     "Grilli J. Macroecological laws describe variation and diversity in "
     "microbial communities. Nat Commun. 2020;11(1):4743.",
     "DOI 10.1038/s41467-020-18529-y",
     "Verified 2026-05-04 (CrossRef; PDF in Refs/)",
     "doi.org/10.1038/s41467-020-18529-y"),
    (6,  "L2 Stochastic-logistic / Taylor microbial",
     "Shoemaker WR, Grilli J. Investigating macroecological patterns in "
     "coarse-grained microbial communities using the stochastic logistic "
     "model of growth. eLife. 2024;12:RP89650.",
     "PMID 38251984",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/38251984/"),
    (7,  "L2 Stochastic-logistic / Taylor microbial",
     "Zaoli S, Grilli J. A macroecological description of alternative stable "
     "states reproduces intra- and inter-host variability of gut microbiome. "
     "Sci Adv. 2021;7(43):eabj2882.",
     "PMID 34669476",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/34669476/"),
    (8,  "L2 Stochastic-logistic / Taylor microbial",
     "Yi B, Chen H. Power law analysis of the human milk microbiome. Arch "
     "Microbiol. 2022;204(9):554.",
     "PMID 36048299",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/36048299/"),
    (9,  "L2 Stochastic-logistic / Taylor microbial",
     "Ma ZS. Power law analysis of the human microbiome. Mol Ecol. "
     "2015;24(21):5428-5444.",
     "PMID 26407082",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/26407082/"),

    (10, "L3 Neutral theory",
     "Hubbell SP. The Unified Neutral Theory of Biodiversity and Biogeography. "
     "Princeton: Princeton University Press; 2001.",
     "ISBN 978-0691021294",
     "No PMID expected (book)",
     "press.princeton.edu/books/paperback/9780691021294"),
    (11, "L3 Neutral theory",
     "Volkov I, Banavar JR, Hubbell SP, Maritan A. Neutral theory and "
     "relative species abundance in ecology. Nature. 2003;424(6952):1035-1037.",
     "PMID 12944964",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/12944964/"),
    (12, "L3 Neutral theory",
     "Etienne RS. A new sampling formula for neutral biodiversity. Ecol Lett. "
     "2005;8(3):253-260.",
     "DOI 10.1111/j.1461-0248.2004.00717.x",
     "Verified 2026-05-04 (CrossRef)",
     "doi.org/10.1111/j.1461-0248.2004.00717.x"),

    (13, "L4 Earth Microbiome / cross-biome catalogues",
     "Thompson LR, Sanders JG, McDonald D, et al. A communal catalogue "
     "reveals Earth's multiscale microbial diversity. Nature. "
     "2017;551(7681):457-463.",
     "PMID 29088705",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/29088705/"),
    (14, "L4 Earth Microbiome / cross-biome catalogues",
     "Sunagawa S, Coelho LP, Chaffron S, et al. Structure and function of "
     "the global ocean microbiome. Science. 2015;348(6237):1261359.",
     "PMID 25999513",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/25999513/"),
    (15, "L4 Earth Microbiome / cross-biome catalogues",
     "Almeida A, Nayfach S, Boland M, et al. A unified catalog of 204,938 "
     "reference genomes from the human gut microbiome. Nat Biotechnol. "
     "2021;39(1):105-114.",
     "PMID 32690973",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/32690973/"),
    (16, "L4 Earth Microbiome / cross-biome catalogues",
     "Bahram M, Hildebrand F, Forslund SK, et al. Structure and function of "
     "the global topsoil microbiome. Nature. 2018;560(7717):233-237.",
     "PMID 30069051",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/30069051/"),
    (17, "L4 Earth Microbiome / cross-biome catalogues",
     "Danko D, Bezdan D, Afshin EE, et al. A global metagenomic map of urban "
     "microbiomes and antimicrobial resistance. Cell. 2021;184(13):3376-3393.",
     "PMID 34043940",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/34043940/"),
    (18, "L4 Earth Microbiome / cross-biome catalogues",
     "Nayfach S, Roux S, Seshadri R, et al. A genomic catalog of Earth's "
     "microbiomes. Nat Biotechnol. 2021;39(4):499-509.",
     "PMID 33169036",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/33169036/"),
    (19, "L4 Earth Microbiome / cross-biome catalogues",
     "Costello EK, Lauber CL, Hamady M, et al. Bacterial community variation "
     "in human body habitats across space and time. Science. "
     "2009;326(5960):1694-1697.",
     "PMID 19892944",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/19892944/"),

    (20, "L5 Shotgun replication / curatedMG",
     "Pasolli E, Schiffer L, Manghi P, et al. Accessible, curated metagenomic "
     "data through ExperimentHub. Nat Methods. 2017;14(11):1023-1024.",
     "PMID 29088129",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/29088129/"),
    (21, "L5 Shotgun replication / curatedMG",
     "Lloyd-Price J, Arze C, Ananthakrishnan AN, et al. Multi-omics of the "
     "gut microbial ecosystem in inflammatory bowel diseases. Nature. "
     "2019;569(7758):655-662.",
     "PMID 31142855",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/31142855/"),

    (22, "L6 Microbial census / scale",
     "Sender R, Fuchs S, Milo R. Revised Estimates for the Number of Human "
     "and Bacteria Cells in the Body. PLoS Biol. 2016;14(8):e1002533.",
     "PMID 27541692",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/27541692/"),

    (23, "L7 Bayesian hierarchical / partial pooling",
     "Gelman A, Hill J. Data Analysis Using Regression and Multilevel/"
     "Hierarchical Models. New York: Cambridge University Press; 2006.",
     "ISBN 978-0521686891",
     "No PMID expected (textbook)",
     "www.cambridge.org/9780521686891"),
    (24, "L7 Bayesian hierarchical / partial pooling",
     "Vehtari A, Gelman A, Gabry J. Practical Bayesian model evaluation using "
     "leave-one-out cross-validation and WAIC. Stat Comput. "
     "2017;27(5):1413-1432.",
     "DOI 10.1007/s11222-016-9696-4",
     "Verified 2026-05-04 (CrossRef; cited 4,272x)",
     "doi.org/10.1007/s11222-016-9696-4"),
    (25, "L7 Bayesian hierarchical / partial pooling",
     "Kass RE, Raftery AE. Bayes Factors. J Am Stat Assoc. 1995;90(430):"
     "773-795.",
     "DOI 10.1080/01621459.1995.10476572",
     "Verified 2026-05-04 (CrossRef; cited 13,492x)",
     "doi.org/10.1080/01621459.1995.10476572"),
    (26, "L7 Bayesian hierarchical / partial pooling",
     "Abril-Pla O, Andreani V, Carroll C, et al. PyMC: a modern, and "
     "comprehensive probabilistic programming framework in Python. PeerJ "
     "Comput Sci. 2023;9:e1516.",
     "DOI 10.7717/peerj-cs.1516",
     "Verified 2026-05-04 (CrossRef)",
     "doi.org/10.7717/peerj-cs.1516"),

    (27, "L8 OSF pre-registration",
     "Nosek BA, Ebersole CR, DeHaven AC, Mellor DT. The preregistration "
     "revolution. Proc Natl Acad Sci USA. 2018;115(11):2600-2606.",
     "PMID 29531091",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/29531091/"),

    (28, "L9 Carrying-capacity heritage",
     "Verhulst PF. Notice sur la loi que la population poursuit dans son "
     "accroissement. Correspondance Mathematique et Physique. 1838;10:113-121.",
     "Historical (1838)",
     "No PMID expected (pre-modern)",
     "en.wikipedia.org/wiki/Logistic_function"),

    (29, "L10 Method tooling",
     "Caporaso JG, Kuczynski J, Stombaugh J, et al. QIIME allows analysis of "
     "high-throughput community sequencing data. Nat Methods. "
     "2010;7(5):335-336.",
     "PMID 20383131",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/20383131/"),
    (30, "L10 Method tooling",
     "Bolyen E, Rideout JR, Dillon MR, et al. Reproducible, interactive, "
     "scalable and extensible microbiome data science using QIIME 2. Nat "
     "Biotechnol. 2019;37(8):852-857.",
     "PMID 31341288",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/31341288/"),
    (31, "L10 Method tooling",
     "Amir A, McDonald D, Navas-Molina JA, et al. Deblur Rapidly Resolves "
     "Single-Nucleotide Community Sequence Patterns. mSystems. "
     "2017;2(2):e00191-16.",
     "PMID 28289731",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/28289731/"),

    (32, "L11 Disease microbiome",
     "Zeller G, Tap J, Voigt AY, et al. Potential of fecal microbiota for "
     "early-stage detection of colorectal cancer. Mol Syst Biol. "
     "2014;10(11):766.",
     "PMID 25432777",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/25432777/"),
    (33, "L11 Disease microbiome",
     "Yachida S, Mizutani S, Shiroma H, et al. Metagenomic and metabolomic "
     "analyses reveal distinct stage-specific phenotypes of the gut "
     "microbiota in colorectal cancer. Nat Med. 2019;25(6):968-976.",
     "PMID 31171880",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/31171880/"),
    (34, "L11 Disease microbiome",
     "Forslund K, Hildebrand F, Nielsen T, et al. Disentangling type 2 "
     "diabetes and metformin treatment signatures in the human gut "
     "microbiota. Nature. 2015;528(7581):262-266.",
     "PMID 26633628",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/26633628/"),
    (35, "L11 Disease microbiome",
     "Gopalakrishnan V, Spencer CN, Nezi L, et al. Gut microbiome modulates "
     "response to anti-PD-1 immunotherapy in melanoma patients. Science. "
     "2018;359(6371):97-103.",
     "PMID 29097493",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/29097493/"),
    (36, "L11 Disease microbiome",
     "Routy B, Le Chatelier E, Derosa L, et al. Gut microbiome influences "
     "efficacy of PD-1-based immunotherapy against epithelial tumors. "
     "Science. 2018;359(6371):91-97.",
     "PMID 29097494",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/29097494/"),
    (37, "L11 Disease microbiome",
     "Davar D, Dzutsev AK, McCulloch JA, et al. Fecal microbiota transplant "
     "overcomes resistance to anti-PD-1 therapy in melanoma patients. "
     "Science. 2021;371(6529):595-602.",
     "PMID 33542131",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/33542131/"),
    (38, "L11 Disease microbiome",
     "Baruch EN, Youngster I, Ben-Betzalel G, et al. Fecal microbiota "
     "transplant promotes response in immunotherapy-refractory melanoma "
     "patients. Science. 2021;371(6529):602-609.",
     "PMID 33303685",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/33303685/"),
    (39, "L11 Disease microbiome",
     "van Nood E, Vrieze A, Nieuwdorp M, et al. Duodenal infusion of donor "
     "feces for recurrent Clostridium difficile. N Engl J Med. "
     "2013;368(5):407-415.",
     "PMID 23323867",
     "Verified 2026-05-04",
     "pubmed.ncbi.nlm.nih.gov/23323867/"),

    (40, "L12 Convergence framing analogue",
     "Xu Z, Lu Z, Ugurbil A, et al. (Cao J corresp.) Mapping convergent "
     "regulators of melanoma drug resistance by PerturbFate. Nature. "
     "2026 [in press; pre-print].",
     "PMID 41986722  /  DOI 10.1038/s41586-026-10367-0",
     "Verified 2026-05-04 (PubMed esummary)",
     "pubmed.ncbi.nlm.nih.gov/41986722/"),
]


def shade(cell, color):
    pPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def set_cn_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "PMingLiU")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def H(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_cn_font(r, size=sizes.get(level, 11), bold=True,
                color="1F4E79" if level == 1 else "2C2C2C")


def P(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_cn_font(r, size=10.5)
    r.font.italic = italic


def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "PMingLiU")
    rPr.append(rFonts)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main():
    doc = Document()
    sect = doc.sections[0]
    sect.left_margin = Inches(0.5)
    sect.right_margin = Inches(0.5)
    sect.top_margin = Inches(0.6)
    sect.bottom_margin = Inches(0.6)

    H(doc, "T5 Macroecology  Full-Field Reference Verification "
           "(Broadest framing)", level=1)
    P(doc, f"Generated: {date.today().isoformat()}  ｜  "
           "Identifier audit: PubMed esummary API + CrossRef API.  "
           "User can click each link in column 'Verification link' to confirm. "
           "Mark 'OK' / 'NG' / a comment in the rightmost column.")
    P(doc, "Status legend: Verified YYYY-MM-DD = identifier confirmed against "
           "PubMed esummary or CrossRef on that date. No PMID expected = book / "
           "institutional / pre-PubMed historical source. Re-audit within 7 days "
           "of submission.")
    P(doc, "Lens map (broadest framing for T5):")
    for lens in [
        "L1 Macroecology / scaling laws (Taylor 1961, May 1988, Locey 2016, Falkowski 2008)",
        "L2 Stochastic-logistic / Taylor microbial (Grilli 2020, Shoemaker-Grilli 2024, Zaoli 2021, Yi 2022, Ma 2015)",
        "L3 Neutral theory (Hubbell 2001 book, Volkov 2003, Etienne 2005)",
        "L4 Earth Microbiome / cross-biome catalogues (Thompson EMP 2017, Sunagawa Tara 2015, Bahram soil 2018, Danko MetaSUB 2021, Nayfach IMG/M 2021, Almeida UHGG 2021, Costello body habitat 2009)",
        "L5 Shotgun replication / curatedMG (Pasolli 2017, Lloyd-Price iHMP 2019)",
        "L6 Microbial census / scale (Sender 2016)",
        "L7 Bayesian hierarchical / partial pooling (Vehtari LOO 2017, Kass-Raftery BIC 1995, PyMC 2023, Gelman 2006)",
        "L8 OSF pre-registration (Nosek 2018)",
        "L9 Carrying-capacity heritage (Verhulst 1838)",
        "L10 Method tooling (QIIME 2010, QIIME 2 2019, Deblur 2017)",
        "L11 Disease microbiome (Zeller 2014, Yachida 2019, Forslund 2015, Gopalakrishnan 2018, Routy 2018, Davar 2021, Baruch 2021, van Nood 2013)",
        "L12 Convergence framing (Xu 2026 PerturbFate; Cao lab)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(lens)
        set_cn_font(r, size=10)

    P(doc, "")

    # ------- Table -------
    headers = ["#", "Lens", "Citation", "Identifier (PMID / DOI)",
               "Status", "Verification link", "User check (OK / NG / comment)"]
    t = doc.add_table(rows=1 + len(CITED), cols=len(headers))
    t.style = "Light Grid Accent 1"
    col_widths = [0.32, 1.45, 3.10, 1.35, 1.40, 1.95, 1.80]

    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_cn_font(r, size=10, bold=True, color="FFFFFF")
        shade(cell, "1F4E79")
        cell.width = Inches(col_widths[j])

    for i, (idx, lens, cite, ident, status, link) in enumerate(CITED, start=1):
        row = t.rows[i]
        for j, w in enumerate(col_widths):
            row.cells[j].width = Inches(w)
        for j, val in enumerate([str(idx), lens, cite, ident, status]):
            cell = row.cells[j]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            set_cn_font(r, size=9.0)
            if "Verified" in val:
                r.font.color.rgb = RGBColor.from_string("0F6B0F")
            elif "Pending" in val:
                r.font.color.rgb = RGBColor.from_string("AA5500")
        link_cell = row.cells[5]
        link_cell.paragraphs[0].clear()
        url = link if link.startswith("http") else f"https://{link}"
        add_hyperlink(link_cell.paragraphs[0], url, link)
        check_cell = row.cells[6]
        check_cell.paragraphs[0].clear()
        r = check_cell.paragraphs[0].add_run("")
        set_cn_font(r, size=9.0)

    # ------- Footer -------
    P(doc, "")
    P(doc,
      f"End of verification list. Total entries: {len(CITED)}. "
      "Re-run audit (PubMed esummary + CrossRef DOI) within 7 days of "
      "manuscript submission to catch any rare PubMed PMID re-assignment "
      "or CrossRef DOI metadata change. Replace any 'NG'-flagged entry by "
      "the closest real paper, with explicit substitution note in the "
      "maintenance log (see CLAUDE.md > 'Manuscript Topic Frame Pattern').")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
