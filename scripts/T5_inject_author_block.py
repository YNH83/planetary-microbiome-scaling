"""
Inject author block + affiliations + corresponding-author block into the
T5 v4.1 manuscript and cover letter v0.2.

Source of author metadata: 04202026 A shared Taylor links biomes.docx
(this fragment was created 2026-04-20 and carries the canonical author
block; it has not been merged into v4.x).

Outputs:
    1. T5_title_intro_methods_results_discussion_package_v4.2_with_authors.docx
       (clone of v4.1; inserts author block + affiliations + corresponding
        authors immediately after the recommended title)
    2. drafts/T5_cover_letter_v0.3.md  (clone of v0.2; replaces the
       "(author block to be finalised)" placeholder with the full block)
    3. T5_cover_letter_v0.3.docx (rebuilt from v0.3 markdown)

ORCIDs (all three confirmed 2026-05-04):
    Yu-Nan Huang       ORCID 0000-0003-1688-1685
    Pen-Hua Su         ORCID 0000-0003-4174-5036
    Chieh-Chen Huang   ORCID 0000-0002-3739-6315
"""
from __future__ import annotations
from copy import deepcopy
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path("/Users/ynh83/Desktop/T5_Macroecology")
SRC_MS = ROOT / "T5_title_intro_methods_results_discussion_package_v4.1_with_references.docx"
OUT_MS = ROOT / "T5_title_intro_methods_results_discussion_package_v4.2_with_authors.docx"

SRC_CL_MD = Path("/Users/ynh83/Desktop/Epi-Protocols/04152026 Microbiome-Epi Protocols/drafts/T5_cover_letter_v0.2.md")
OUT_CL_MD = Path("/Users/ynh83/Desktop/Epi-Protocols/04152026 Microbiome-Epi Protocols/drafts/T5_cover_letter_v0.3.md")
OUT_CL_DOCX = ROOT / "T5_cover_letter_v0.3.docx"

AUTHOR_LINE = "Yu-Nan Huang [1,2,3], Pen-Hua Su [2,3,*], Chieh-Chen Huang [1,*]"
AFFILIATIONS = (
    "1 Department of Life Science, National Chung Hsing University, Taichung, Taiwan; "
    "2 Division of Genetics and Metabolism, Department of Pediatrics, Chung Shan Medical "
    "University Hospital, Taichung, Taiwan; "
    "3 School of Medicine, Chung Shan Medical University, Taichung, Taiwan."
)
CORRESPONDENCE = (
    "Correspondence to: Pen-Hua Su, Tel: +886 4 2473 9595 (Ext. 21707), "
    "email: ninaphsu@gmail.com (ORCID 0000-0003-4174-5036); and Chieh-Chen Huang, "
    "Tel: +886 4 2284 0416 (Ext. 405), email: cchuang@dragon.nchu.edu.tw "
    "(ORCID 0000-0002-3739-6315)."
)
ORCID_BLOCK = (
    "ORCIDs: Yu-Nan Huang 0000-0003-1688-1685; Pen-Hua Su 0000-0003-4174-5036; "
    "Chieh-Chen Huang 0000-0002-3739-6315."
)


def add_para(doc, text, *, bold=False, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def inject_into_manuscript():
    doc = Document(str(SRC_MS))
    body = doc.element.body
    # Find the "Recommended title" anchor and the next non-empty paragraph
    # which is the title text. Insert author block right after the title.
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Convergent Taylor scaling links planetary microbiomes through a habitat-modulated carrying-capacity axis":
            target_idx = i
            break
    if target_idx is None:
        raise RuntimeError("Title paragraph not found in v4.1")

    # Build new paragraphs in a stub doc, then insert XML
    stub = Document()
    stub.add_paragraph()  # spacer
    p_auth = stub.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_auth.add_run(AUTHOR_LINE)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11.5)

    p_aff = stub.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_aff.add_run(AFFILIATIONS)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    p_corr = stub.add_paragraph()
    p_corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_corr.add_run(CORRESPONDENCE)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    p_orcid = stub.add_paragraph()
    p_orcid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_orcid.add_run(ORCID_BLOCK)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.5)

    stub.add_paragraph()  # trailing spacer

    # Insert each new paragraph right after target_idx
    target_p = doc.paragraphs[target_idx]
    target_xml = target_p._element
    for new_p in stub.paragraphs:
        target_xml.addnext(deepcopy(new_p._element))
        target_xml = target_xml.getnext()

    doc.save(str(OUT_MS))
    print(f"Wrote: {OUT_MS}  size={OUT_MS.stat().st_size:,} bytes")


def update_cover_letter():
    md = SRC_CL_MD.read_text()

    # Replace closing block
    new_close = (
        f"Sincerely,\n\n"
        f"{AUTHOR_LINE}\n\n"
        f"Affiliations:\n"
        f"1. Department of Life Science, National Chung Hsing University, "
        f"Taichung, Taiwan\n"
        f"2. Division of Genetics and Metabolism, Department of Pediatrics, "
        f"Chung Shan Medical University Hospital, Taichung, Taiwan\n"
        f"3. School of Medicine, Chung Shan Medical University, Taichung, Taiwan\n\n"
        f"Corresponding authors:\n"
        f"- Pen-Hua Su, Tel: +886 4 2473 9595 (Ext. 21707), email: "
        f"ninaphsu@gmail.com (ORCID 0000-0003-4174-5036)\n"
        f"- Chieh-Chen Huang, Tel: +886 4 2284 0416 (Ext. 405), email: "
        f"cchuang@dragon.nchu.edu.tw (ORCID 0000-0002-3739-6315)\n\n"
        f"ORCIDs: Yu-Nan Huang 0000-0003-1688-1685; Pen-Hua Su 0000-0003-4174-5036; "
        f"Chieh-Chen Huang 0000-0002-3739-6315.\n"
    )
    md_new = md.replace("Sincerely,\n\n(author block to be finalised)\n", new_close)
    if md_new == md:
        # fallback regex
        md_new = re.sub(
            r"Sincerely,\s*\n\s*\(author block to be finalised\)\s*\n?",
            new_close,
            md,
        )
    if md_new == md:
        raise RuntimeError("Could not locate placeholder in cover letter v0.2")

    # Bump header date and Re: stays same
    md_new = md_new.replace("**Date**: 2026-05-04",
                            "**Date**: 2026-05-04 (v0.3 author block injected)")

    OUT_CL_MD.write_text(md_new)
    print(f"Wrote: {OUT_CL_MD}  size={OUT_CL_MD.stat().st_size:,} bytes")

    # Build docx
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    bold_re = re.compile(r"\*\*(.+?)\*\*")

    def add_text(text):
        para = d.add_paragraph()
        pos = 0
        for m in bold_re.finditer(text):
            if m.start() > pos:
                para.add_run(text[pos:m.start()])
            r = para.add_run(m.group(1))
            r.bold = True
            pos = m.end()
        if pos < len(text):
            para.add_run(text[pos:])

    for line in md_new.splitlines():
        if not line.strip():
            d.add_paragraph()
            continue
        if line.startswith("# "):
            d.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            d.add_heading(line[3:], level=2)
        elif re.match(r"^\d+\.\s", line):
            d.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            d.add_paragraph(line[2:], style="List Bullet")
        else:
            add_text(line)
    d.save(str(OUT_CL_DOCX))
    print(f"Wrote: {OUT_CL_DOCX}  size={OUT_CL_DOCX.stat().st_size:,} bytes")


if __name__ == "__main__":
    inject_into_manuscript()
    update_cover_letter()
