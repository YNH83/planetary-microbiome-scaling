"""
Bring the three top-level Chinese docx files into the 2026-05-04 evening
sync standard:

    1. T5_K承載力疾病量化_ZH.docx  (clinical translation / K leverage chapter)
    2. T5_完整研究說明_ZH.docx     (full study explanation)
    3. T5_概念與圖表詳解_2026-05-04.docx  (concept walkthrough)

Each file gets a "投稿同步資訊 (2026-05-04 evening sync)" frontmatter block
inserted immediately after its own first paragraph. The block carries:

    - convergence English title
    - author block (Yu-Nan Huang, Pen-Hua Su, Chieh-Chen Huang)
    - affiliations (NCHU 1; CSMU Hospital 2; CSMU School of Medicine 3)
    - corresponding-author lines (Pen-Hua Su, Chieh-Chen Huang)
    - ORCIDs
    - pointer to canonical v4.2 manuscript and v0.3 cover letter
    - explicit role of this file inside the 2026-05-04 package

Idempotent: the script detects an existing sync banner via the marker
"[投稿同步資訊 2026-05-04 evening sync]" and replaces it in place,
so it can be re-run after future evening syncs.

Outputs in place; PostToolUse hook strips em / en dashes if any sneak in.
"""
from __future__ import annotations
from copy import deepcopy
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/ynh83/Desktop/T5_Macroecology")

TARGETS = [
    (ROOT / "T5_K承載力疾病量化_ZH.docx",
     "本檔案在 2026-05-04 evening sync 中的角色: 槓桿節點 (D) 的臨床轉譯展開層。"
     "提供 IBD 復發、大腸癌、免疫治療、藥物動力學等場景下「K 承載力作為可計算偏差度量」"
     "的詳細推導與應用。"),
    (ROOT / "T5_完整研究說明_ZH.docx",
     "本檔案在 2026-05-04 evening sync 中的角色: 整體敘事層。提供 T5 收斂研究的完整中文展開, "
     "涵蓋背景, 方法, 結果, 討論, 與前置註冊保護機制。"),
    (ROOT / "T5_概念與圖表詳解_2026-05-04.docx",
     "本檔案在 2026-05-04 evening sync 中的角色: 指標與圖表詳解層。提供 100 講統計指標, "
     "5 張主圖逐欄拆解, 13 張 supplementary 圖示, 與審稿人攻擊面 Q&A。"),
]

MARKER = "[投稿同步資訊 2026-05-04 evening sync]"
CONVERGENCE_EN = (
    "Convergent Taylor scaling links planetary microbiomes through a "
    "habitat-modulated carrying-capacity axis"
)
CONVERGENCE_ZH = (
    "收斂的 Taylor 尺度將行星尺度微生物體連結於一條棲地調控的承載量軸"
)

AUTHOR_LINE = "Yu-Nan Huang [1,2,3], Pen-Hua Su [2,3,*], Chieh-Chen Huang [1,*]"
AFFILIATIONS = (
    "1. Department of Life Science, National Chung Hsing University, "
    "Taichung, Taiwan",
    "2. Division of Genetics and Metabolism, Department of Pediatrics, "
    "Chung Shan Medical University Hospital, Taichung, Taiwan",
    "3. School of Medicine, Chung Shan Medical University, Taichung, Taiwan",
)
CORRESPONDENCE = (
    "Pen-Hua Su, Tel: +886 4 2473 9595 (Ext. 21707), email: ninaphsu@gmail.com "
    "(ORCID 0000-0003-4174-5036)",
    "Chieh-Chen Huang, Tel: +886 4 2284 0416 (Ext. 405), email: "
    "cchuang@dragon.nchu.edu.tw (ORCID 0000-0002-3739-6315)",
)
YN_ORCID = "Yu-Nan Huang ORCID 0000-0003-1688-1685"

CANONICAL_DOCS = (
    "T5_title_intro_methods_results_discussion_package_v4.2_with_authors.docx "
    "(latest manuscript)",
    "T5_cover_letter_v0.3.docx (latest cover letter)",
    "T5_References_FullField_Verification_2026-05-04.docx (40-entry reference "
    "verification across 12 lenses)",
    "SUBMISSION_CHECKLIST.md (10-section go / no-go at package root)",
    "drafts/T5_post_hoc_framing_note_2026-05-04.md (OSF preregistration "
    "integrity protection)",
)


def make_run(paragraph, text, *, bold=False, italic=False, size=10.5,
             east_asia=True, color=None):
    r = paragraph.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if east_asia:
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "PMingLiU")
    return r


def build_banner_paragraphs(role_note, stub):
    """Build sync banner inside a stub Document and return a list of paragraph
    XML elements ready to be inserted into the target doc."""
    paragraphs = []

    p = stub.add_paragraph()
    make_run(p, MARKER, bold=True, size=11, color="0E4D92")
    paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "Recommended title (English): ", bold=True)
    make_run(p, CONVERGENCE_EN)
    paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "建議標題 (中文): ", bold=True)
    make_run(p, CONVERGENCE_ZH)
    paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "Authors: ", bold=True)
    make_run(p, AUTHOR_LINE)
    paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "Affiliations:", bold=True)
    paragraphs.append(p)
    for aff in AFFILIATIONS:
        p = stub.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        make_run(p, aff)
        paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "Corresponding authors:", bold=True)
    paragraphs.append(p)
    for c in CORRESPONDENCE:
        p = stub.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        make_run(p, c)
        paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "ORCID: ", bold=True)
    make_run(p, YN_ORCID + "; Pen-Hua Su 0000-0003-4174-5036; "
             "Chieh-Chen Huang 0000-0002-3739-6315.")
    paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "Companion documents (canonical, 2026-05-04 evening sync):",
             bold=True)
    paragraphs.append(p)
    for c in CANONICAL_DOCS:
        p = stub.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        make_run(p, "- " + c)
        paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "本檔案角色: ", bold=True)
    make_run(p, role_note)
    paragraphs.append(p)

    p = stub.add_paragraph()
    make_run(p, "OSF preregistration integrity: ", bold=True)
    make_run(p,
             "T5_OSF_preregistration_v0.2.{md,docx} 與 T5_OSF_Step{4,5,6,7}_*.docx "
             "保持原樣未動; H1 至 H7 預註冊假設與門檻不變。所有 2026-05-04 變更皆為編輯重定位, "
             "非分析重做。")
    paragraphs.append(p)

    p = stub.add_paragraph()  # trailing spacer
    paragraphs.append(p)

    return paragraphs


def remove_existing_banner(doc):
    """Remove any prior MARKER block and trailing paragraphs up to the next
    non-banner content. We mark the banner span by walking from the marker
    paragraph until we hit an empty paragraph (the trailing spacer)."""
    body = doc.element.body
    paras = list(doc.paragraphs)
    start = None
    for i, p in enumerate(paras):
        if MARKER in p.text:
            start = i
            break
    if start is None:
        return False
    # Banner ends at the next paragraph whose text starts with "本檔案角色"
    # plus the OSF integrity paragraph plus a trailing empty paragraph.
    end = start
    seen_role = False
    seen_osf = False
    for j in range(start, min(start + 30, len(paras))):
        text = paras[j].text
        if text.startswith("本檔案角色"):
            seen_role = True
        elif text.startswith("OSF preregistration integrity"):
            seen_osf = True
            end = j
        elif seen_role and seen_osf and text.strip() == "":
            end = j
            break
    # Remove paragraphs start..end inclusive
    for p in paras[start:end + 1]:
        p._element.getparent().remove(p._element)
    return True


def sync_one(target_path, role_note):
    doc = Document(str(target_path))
    removed = remove_existing_banner(doc)

    # Use a stub document to build paragraphs cleanly
    stub = Document()
    new_paras = build_banner_paragraphs(role_note, stub)

    # Insert after the first non-empty paragraph (the doc's own title)
    target_idx = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            target_idx = i
            break
    anchor_xml = doc.paragraphs[target_idx]._element
    for new_p in new_paras:
        anchor_xml.addnext(deepcopy(new_p._element))
        anchor_xml = anchor_xml.getnext()

    doc.save(str(target_path))
    return removed


if __name__ == "__main__":
    for target_path, role_note in TARGETS:
        before_size = target_path.stat().st_size
        replaced = sync_one(target_path, role_note)
        after_size = target_path.stat().st_size
        verb = "REPLACED" if replaced else "INSERTED"
        delta = after_size - before_size
        sign = "+" if delta >= 0 else ""
        print(f"{verb:8s}  {target_path.name}  "
              f"{before_size:,} -> {after_size:,} bytes  ({sign}{delta:,})")
    print("\nDone.")
