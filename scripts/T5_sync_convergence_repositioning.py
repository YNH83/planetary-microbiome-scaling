"""
T5 master sync: propagate the 2026-05-04 convergence repositioning across
every live document in the project.

This script captures the 2026-05-04 morning sync (steps 1 to 13). The
2026-05-04 evening sync added five additional outputs that are produced
by separate dedicated scripts (listed below) and do not need to be
re-run from this master script:

    - T5_References_FullField_Verification_2026-05-04.docx
        produced by: scripts/T5_build_full_lit_verification.py
    - T5_..._v4.1_with_references.docx
        produced by: scripts/T5_append_references_to_v4.py
    - T5_..._v4.2_with_authors.docx
    - T5_cover_letter_v0.3.{md,docx}
        produced by: scripts/T5_inject_author_block.py
    - T5_cover_letter_v0.2.{md,docx}
        produced by: hand (md) + same inject script (docx)

What this script does (non-destructive; preserves OSF preregistration intact):
    1. drafts/T5_manuscript_v0.2_EN.md      title + Fig.5 path / caption
    2. drafts/T5_manuscript_v0.2_ZH.md      title + Fig.5 path / caption
    3. drafts/T5_cover_letter_v0.1.md       Re: line
    4. T5_..._v4_with_figures.docx          title + repositioning note (in place)
    5. T5_完整研究說明_ZH.docx               title only (in place)
    6. T5_K承載力疾病量化_ZH.docx            append cross-link note (in place)
    7. figures/renamed/                     copy graphical abstract + leverage Fig.5
    8. figures/renamed/FIGURE_DATA_PROVENANCE.md  add 2 rows
    9. README.md                            add convergence framing block
   10. wiki/log.md                          append WIKI entry
   11. wiki/t5_topic_frame.md               prepend convergence frame
   12. wiki/submissions.md                  add T5 repositioning sub-bullet
   13. drafts/T5_post_hoc_framing_note_2026-05-04.md   NEW (OSF protection)

What this DOES NOT do (intentional):
    - drafts/T5_OSF_preregistration_*       preregistration is integrity-locked
    - T5_OSF_Step{4,5,6,7}_*.docx           same
    - T5_References_Verification_*.docx     locked verification artifact
    - T5_..._v3.docx                        v3 is locked; v3.1 already exists
    - drafts/T5_manuscript_v0.1_*           historical
    - T5_..._v4.1_with_references.docx      see scripts/T5_append_references_to_v4.py
    - T5_..._v4.2_with_authors.docx         see scripts/T5_inject_author_block.py
    - T5_cover_letter_v0.{2,3}.docx         see hand draft + inject script
"""
from __future__ import annotations
import shutil
import re
from pathlib import Path
from datetime import date
from docx import Document

ROOT = Path("/Users/ynh83/Desktop/T5_Macroecology")
WIKI = ROOT / "wiki"
DRAFTS = ROOT / "drafts"
FIGS = ROOT / "figures"
RENAMED = FIGS / "renamed"

NEW_TITLE_EN = (
    "Convergent Taylor scaling links planetary microbiomes through a "
    "habitat-modulated carrying-capacity axis"
)
NEW_TITLE_ZH = (
    "收斂的 Taylor 尺度將行星尺度微生物體連結於一條棲地調控的承載量軸 "
    f"({NEW_TITLE_EN})"
)
OLD_EN_LEAD = ("One macroecological law governs gut, soil, ocean, and air "
               "microbiomes: Bayesian hierarchical universality with invariant "
               "exponent and habitat-modulated carrying capacity")
OLD_ZH_LEAD = ("一條宏觀生態定律統攝腸道、土壤、海洋與空氣微生物體：具不變指數"
               "與生境調控承載量的貝氏階層通用性")
OLD_DOCX_TITLE = ("Bayesian hierarchical universality of Taylor's law across "
                  "planetary microbiomes with habitat-modulated carrying "
                  "capacity")

REPOSITIONING_NOTE = (
    "Editorial repositioning (2026-05-04). Title and opening lines have been "
    "rewritten to lead with the convergence framing: different perturbations "
    "(15 EMPO-3 biomes) -> common state (shared Taylor backbone, beta near 2) "
    "-> common program (stochastic-logistic + Gamma AFD) -> leverage node "
    "(habitat / disease / time act on K, not on beta). The framing mirrors "
    "Cao et al. Nature 2026 (PerturbFate). The Bayesian-hierarchical "
    "methodology is retained verbatim in Methods; preregistered hypotheses "
    "and OSF documents are not retroactively modified. See "
    "T5_graphical_abstract.png (graphical abstract v1) and "
    "T5_fig5_leverage.png (Fig.5 \"K is the leverage point\") for matched "
    "figure-level repositioning."
)
ZH_REPOSITIONING_NOTE = (
    "編輯重定位 (2026-05-04)。標題與開場已改為以收斂框架開頭: "
    "不同擾動 (15 個 EMPO-3 棲地) → 共同狀態 (共享 Taylor 主軸, beta 接近 2) "
    "→ 共同程序 (stochastic-logistic + Gamma AFD) → 槓桿節點 "
    "(棲地 / 疾病 / 時間進入 K, 而非 beta)。此框架對應 Cao et al. Nature 2026 "
    "(PerturbFate) 的四步敘事。Bayesian hierarchical 方法逐字保留於 Methods; "
    "預註冊假設與 OSF 文件不做回溯性修改。圖見 "
    "T5_graphical_abstract.png 與 T5_fig5_leverage.png。"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def replace_in_docx(path: Path, edits: list[tuple[str, str]],
                     insert_after_starts_with: tuple[str, str] | None = None,
                     italic_insert: bool = True) -> list[str]:
    doc = Document(path)
    log = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        for old, new in edits:
            if old and old in text:
                new_text = text.replace(old, new)
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    p.add_run(new_text)
                log.append(f"  para {i}: replaced '{old[:50]}...'")
                break
    if insert_after_starts_with is not None:
        anchor, note = insert_after_starts_with
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(anchor) and not any(
                q.text.startswith(note[:30]) for q in doc.paragraphs):
                new_para = p.insert_paragraph_before(note)
                if italic_insert:
                    for run in new_para.runs:
                        run.italic = True
                log.append(f"  para {i}: inserted repositioning note above")
                break
    doc.save(path)
    return log


def replace_in_text(path: Path, edits: list[tuple[str, str]],
                     header_prepend: str = "") -> list[str]:
    body = path.read_text(encoding="utf-8")
    log = []
    for old, new in edits:
        if old in body:
            body = body.replace(old, new)
            log.append(f"  replaced '{old[:60]}...'")
    if header_prepend and header_prepend not in body:
        body = header_prepend + body
        log.append("  prepended header note")
    path.write_text(body, encoding="utf-8")
    return log


# ---------------------------------------------------------------------------
# Sync steps
# ---------------------------------------------------------------------------

def step1_manuscript_en():
    p = DRAFTS / "T5_manuscript_v0.2_EN.md"
    edits = [
        ("# " + OLD_EN_LEAD,
         "# " + NEW_TITLE_EN
         + "\n\n*v0.2.1 repositioning (2026-05-04): title leads with "
           "convergence framing; methodology unchanged. See "
           "`figures/renamed/Figure_GA_graphical_abstract.png` and "
           "`figures/renamed/Figure_5_leverage.png`. The original v0.2 title "
           "was: " + OLD_EN_LEAD + "*"),
        ("`figures/renamed/Figure_5_k_distribution.png`",
         "`figures/renamed/Figure_5_leverage.png` (was "
         "`Figure_5_k_distribution.png` in v0.2; rebuilt 2026-05-04 with 4 "
         "panels: per-biome K ridge, beta forest, IBDMDB UC/CD K-shift, "
         "108 longitudinal subjects)"),
        ("Figure 5: per-biome carrying-capacity K distribution, beta invariance",
         "Figure 5: K is the leverage point (4-panel rebuild) ,  per-biome K "
         "ridge, beta invariance forest, IBDMDB disease K-shift, 108 IBD "
         "subjects' beta invariance"),
    ]
    return p, replace_in_docx_or_text(p, edits)


def replace_in_docx_or_text(p, edits):
    if p.suffix == ".docx":
        return replace_in_docx(p, edits)
    return replace_in_text(p, edits)


def step2_manuscript_zh():
    p = DRAFTS / "T5_manuscript_v0.2_ZH.md"
    new_zh_h1 = (
        "# " + NEW_TITLE_ZH + "\n\n"
        "*v0.2.1 重定位 (2026-05-04): 標題改為以收斂框架開頭, 方法學保留不變。"
        "見 `figures/renamed/Figure_GA_graphical_abstract.png` 與 "
        "`figures/renamed/Figure_5_leverage.png`。原 v0.2 標題: "
        + OLD_ZH_LEAD + "*"
    )
    edits = [
        ("# " + OLD_ZH_LEAD + " (Bayesian hierarchical universality with "
         "invariant exponent and habitat-modulated carrying capacity)",
         new_zh_h1),
        ("`figures/renamed/Figure_5_k_distribution.png`",
         "`figures/renamed/Figure_5_leverage.png` (取代 v0.2 的 "
         "`Figure_5_k_distribution.png`; 2026-05-04 改畫為 4 panel: "
         "per-biome K ridge, beta forest, IBDMDB UC/CD K-shift, "
         "108 longitudinal subjects)"),
        ("Figure 5：各生境承載量 K 分布與 beta 不變性",
         "Figure 5: K is the leverage point (4 panel 重建) - 各生境 K ridge, "
         "beta invariance forest, IBDMDB 疾病 K-shift, 108 IBD 受試者 beta 不變"),
    ]
    return p, replace_in_text(p, edits)


def step3_cover_letter():
    p = DRAFTS / "T5_cover_letter_v0.1.md"
    edits = [
        (("**Re**: \"One macroecological law governs gut, soil, ocean, and "
          "air microbiomes: Bayesian hierarchical universality with invariant "
          "exponent and habitat-modulated carrying capacity\""),
         "**Re**: \"" + NEW_TITLE_EN + "\" "
         "*(v0.1.1 repositioning 2026-05-04; methodology preserved)*"),
    ]
    return p, replace_in_text(p, edits)


def step4_v4_docx():
    p = ROOT / "T5_title_intro_methods_results_discussion_package_v4_with_figures.docx"
    edits = [
        (OLD_DOCX_TITLE, NEW_TITLE_EN),
        ("Recommended wording for the central claim: The data support "
         "hierarchical universality of Taylor's law across planetary "
         "microbiomes: a shared central exponent with habitat-modulated "
         "intercept and carrying-capacity structure.",
         "Recommended wording for the central claim: Planetary microbiomes "
         "share a stochastic-logistic Taylor backbone; habitat enters through "
         "K (the carrying-capacity / intercept axis), not through beta. "
         "Strict complete pooling is rejected; hierarchical universality is "
         "supported."),
        ("One-sentence Results headline: Cross-biome abundance moments "
         "collapse onto a shared Taylor-scaling backbone, while habitat "
         "identity is expressed mainly through intercept structure rather "
         "than slope collapse.",
         "One-sentence Results headline: Cross-biome abundance moments "
         "converge onto a single Taylor-scaling backbone, while habitat "
         "identity is expressed mainly through intercept structure rather "
         "than slope collapse."),
        ("One-sentence Discussion headline: Host-associated and free-living "
         "microbiomes are best understood as structured deviations within a "
         "common macroecological regime, not as separate quantitative "
         "kingdoms.",
         "One-sentence Discussion headline: Host-associated and free-living "
         "microbiomes are best understood as structured deviations within a "
         "common macroecological regime: a shared stochastic-logistic "
         "backbone, with habitat, disease, and time entering through the "
         "carrying-capacity axis rather than through the scaling exponent."),
        ("Figure 5. Hierarchical universality and habitat-modulated carrying "
         "capacity.",
         "Figure 5. K is the leverage point: habitat, disease, and time act "
         "on intercept alpha (approximately log K), not on beta. Four "
         "panels: (a) per-biome K ridge density, (b) beta invariance forest "
         "with global 95% HDI band, (c) HMP IBDMDB K-shift between control / "
         "UC / CD, (d) beta invariance across 108 longitudinal IBD "
         "subjects."),
    ]
    return p, replace_in_docx(
        p, edits,
        insert_after_starts_with=("Editorial recommendation.",
                                   REPOSITIONING_NOTE))


def step5_zh_full_docx():
    p = ROOT / "T5_完整研究說明_ZH.docx"
    edits = [
        ("One macroecological law governs gut, soil, ocean, and air "
         "microbiomes: Bayesian hierarchical universality with invariant "
         "exponent and habitat-modulated carrying capacity",
         NEW_TITLE_EN),
        ("一個宏觀生態學法則統治腸道、土壤、海洋與空氣微生物群落：",
         "收斂的 Taylor 尺度將行星尺度微生物體連結於："),
        ("貝氏階層性的普適指數與棲息地調控的承載力",
         "棲地調控的承載量 (K) 軸"),
    ]
    return p, replace_in_docx(p, edits)


def step6_zh_k_docx():
    p = ROOT / "T5_K承載力疾病量化_ZH.docx"
    note_par = (
        "編輯重定位 (2026-05-04 cross-link)。本文件 (Section 2.5 K 承載力的"
        "詳細展開) 在收斂框架下的角色是「槓桿節點 (D)」: 棲地、疾病、時間皆"
        "透過 K 進入 Taylor 定律, 而非透過 beta。對應主稿件已重定位為「收斂的"
        " Taylor 尺度將行星尺度微生物體連結於一條棲地調控的承載量軸」, 並由 "
        "T5_graphical_abstract.png (上層四節點) 與 T5_fig5_leverage.png "
        "(panel a per-biome K ridge + panel c IBDMDB K-shift) 視覺化。"
    )
    doc = Document(p)
    log = []
    if not any(q.text.startswith("編輯重定位 (2026-05-04 cross-link)")
               for q in doc.paragraphs):
        new_para = doc.paragraphs[1].insert_paragraph_before(note_par)
        for run in new_para.runs:
            run.italic = True
        doc.save(p)
        log.append("  inserted repositioning cross-link near top")
    return p, log


def step7_copy_figures():
    src_ga = FIGS / "T5_graphical_abstract.png"
    src_f5 = FIGS / "T5_fig5_leverage.png"
    dst_ga = RENAMED / "Figure_GA_graphical_abstract.png"
    dst_f5 = RENAMED / "Figure_5_leverage.png"
    log = []
    for s, d in [(src_ga, dst_ga), (src_f5, dst_f5)]:
        if s.exists():
            shutil.copy(s, d)
            log.append(f"  copied {s.name} -> {d.name}")
    return RENAMED, log


def step8_provenance_md():
    p = RENAMED / "FIGURE_DATA_PROVENANCE.md"
    body = p.read_text(encoding="utf-8")
    if "Figure_5_leverage.png" not in body:
        addendum = (
            "\n\n---\n\n"
            "## 2026-05-04 Repositioning addendum\n\n"
            "| Fig | Output PNG | Raw data (public) | Intermediate tables | "
            "Producing script |\n"
            "|---|---|---|---|---|\n"
            "| **Fig 5 v2** | `figures/renamed/Figure_5_leverage.png` | "
            "EMP 90 bp BIOM (per-taxon moments) + curatedMG HMP_2019_ibdmdb "
            "(disease state) + iHMP IBDMDB longitudinal | "
            "`results_csv/T5_empo3_real_moments.csv`, "
            "`results_csv/T5_empo3_real_taylor.csv`, "
            "`results_csv/T5_disease_afd.csv`, "
            "`results_csv/T5_disease_detection_results.csv`, "
            "`results_csv/T5_longitudinal_per_subject.csv`, "
            "`results_csv/T5_longitudinal_results.csv` | "
            "`scripts/T5_fig5_leverage.py` |\n"
            "| **Graphical abstract** | "
            "`figures/renamed/Figure_GA_graphical_abstract.png` | "
            "(schematic; no raw data) | n/a | "
            "`scripts/T5_graphical_abstract.py` |\n\n"
            "Both figures replace the earlier `Figure_5_k_distribution.png` "
            "as the main-text K / leverage panel. The earlier file is "
            "retained for provenance.\n"
        )
        p.write_text(body + addendum, encoding="utf-8")
        return p, ["  appended 2 rows to FIGURE_DATA_PROVENANCE.md"]
    return p, []


def step9_readme():
    p = ROOT / "README.md"
    body = p.read_text(encoding="utf-8")
    log = []
    if "## 2026-05-04 Convergence repositioning" not in body:
        block = (
            "\n## 2026-05-04 Convergence repositioning\n\n"
            "Title and figure-level framing repositioned to lead with the "
            "four-step convergence narrative (Cao et al. Nature 2026 "
            "PerturbFate analogue):\n\n"
            "    different perturbations (15 EMPO-3 biomes)\n"
            "    -> common state (shared Taylor backbone, beta near 2)\n"
            "    -> common program (stochastic-logistic + Gamma AFD)\n"
            "    -> leverage node (habitat / disease / time act on K, not beta)\n\n"
            "New deliverables:\n\n"
            "- `figures/T5_graphical_abstract.png` and "
            "`figures/renamed/Figure_GA_graphical_abstract.png` "
            "(graphical abstract v1; 16:9, 4-node arc + 3 evidence rails + 2 "
            "deliverables)\n"
            "- `figures/T5_fig5_leverage.png` and "
            "`figures/renamed/Figure_5_leverage.png` (Fig.5 rebuild with 4 "
            "panels: per-biome K ridge / beta forest / IBDMDB K-shift / "
            "108 longitudinal subjects)\n"
            "- `T5_title_intro_methods_results_discussion_package_v3.1_"
            "convergence.docx` (v3 manuscript with title rewritten)\n"
            "- `drafts/T5_post_hoc_framing_note_2026-05-04.md` "
            "(post-hoc framing note; OSF preregistration preserved unchanged)\n\n"
            "Recommended title: "
            f"\"{NEW_TITLE_EN}\".\n\n"
            "Build scripts:\n\n"
            "- `scripts/T5_fig5_leverage.py`\n"
            "- `scripts/T5_graphical_abstract.py`\n"
            "- `scripts/T5_rewrite_title_abstract.py`\n"
            "- `scripts/T5_sync_convergence_repositioning.py` (this master "
            "sync script)\n"
        )
        p.write_text(body + block, encoding="utf-8")
        log.append("  appended convergence-repositioning block")
    return p, log


def step10_log_md():
    p = WIKI / "log.md"
    body = p.read_text(encoding="utf-8")
    today = date.today().isoformat()
    entry = (
        f"\n**{today} WIKI / DRAFT** T5 convergence repositioning. "
        "Title rewritten to lead with \"Convergent Taylor scaling links "
        "planetary microbiomes through a habitat-modulated carrying-capacity "
        "axis\" (was: Bayesian hierarchical universality...). Fig.5 rebuilt "
        "as \"K is the leverage point\" with 4 panels (per-biome K ridge, "
        "beta forest, IBDMDB UC/CD K-shift, 108 longitudinal subjects). "
        "Graphical abstract v1 added (16:9, 4-node Cao-style arc + 3 evidence "
        "rails + 2 deliverables). Synced files: drafts/T5_manuscript_v0.2_"
        "{EN,ZH}.md, drafts/T5_cover_letter_v0.1.md, v3.1 + v4 docx, ZH "
        "long-form docx, K承載力 docx, README.md, FIGURE_DATA_PROVENANCE.md, "
        "wiki/t5_topic_frame.md, wiki/submissions.md. OSF preregistration "
        "(Step 4-7 docx + drafts/T5_OSF_preregistration_v0.2.{md,docx}) NOT "
        "modified. Scripts: T5_fig5_leverage.py, T5_graphical_abstract.py, "
        "T5_rewrite_title_abstract.py, T5_sync_convergence_repositioning.py.\n"
    )
    if "T5 convergence repositioning" not in body:
        p.write_text(body + entry, encoding="utf-8")
        return p, ["  appended log entry"]
    return p, []


def step11_topic_frame():
    p = WIKI / "t5_topic_frame.md"
    body = p.read_text(encoding="utf-8")
    log = []
    if "Convergence repositioning (2026-05-04)" not in body:
        prepend = (
            "## Convergence repositioning (2026-05-04)\n\n"
            "T5 has been repositioned under the Cao et al. Nature 2026 "
            "(PerturbFate) four-step convergence framework: different "
            "perturbations (15 EMPO-3 biomes) -> common state (shared Taylor "
            "backbone, beta near 2) -> common program (stochastic-logistic + "
            "Gamma AFD) -> leverage node (K, the carrying-capacity / "
            "intercept axis).\n\n"
            "Recommended title: " + NEW_TITLE_EN + "\n\n"
            "Recommended figure architecture: "
            "Figure_GA_graphical_abstract.png (graphical abstract); "
            "Figure_2_universal_collapse.png (common state); "
            "Figure_3_afd_comparison.png + Figure_4_bic_and_hubbell.png "
            "(common program); Figure_5_leverage.png (leverage node).\n\n"
            "OSF preregistration (Step 4-7) is preserved unchanged; all "
            "preregistered hypotheses retain their original wording.\n\n"
            "---\n\n"
        )
        p.write_text(prepend + body, encoding="utf-8")
        log.append("  prepended convergence repositioning block")
    return p, log


def step12_submissions_md():
    p = WIKI / "submissions.md"
    body = p.read_text(encoding="utf-8")
    log = []
    sentinel = "T5 v0.2.1 convergence repositioning (2026-05-04)"
    if sentinel not in body:
        # insert after the 2026-04-20 multi-paper block, before "## Active topics"
        marker = "## Active topics"
        if marker in body:
            note = (
                "\n**" + sentinel + "**: Title rewritten to "
                "\"" + NEW_TITLE_EN + "\". Fig.5 rebuilt as 4-panel "
                "\"K is the leverage point\" (per-biome K ridge / beta "
                "forest / IBDMDB UC/CD K-shift / 108 longitudinal subjects). "
                "Graphical abstract v1 added. Scripts: "
                "`scripts/T5_fig5_leverage.py`, "
                "`scripts/T5_graphical_abstract.py`, "
                "`scripts/T5_sync_convergence_repositioning.py`. "
                "OSF preregistration NOT modified.\n\n"
            )
            body = body.replace(marker, note + marker)
            p.write_text(body, encoding="utf-8")
            log.append("  inserted T5 convergence-repositioning sub-bullet")
    return p, log


def step13_post_hoc_note():
    p = DRAFTS / "T5_post_hoc_framing_note_2026-05-04.md"
    if not p.exists():
        body = (
            "# T5 post-hoc framing note (2026-05-04)\n\n"
            "**Purpose**: document the editorial repositioning of the T5 "
            "manuscript without modifying any preregistered material.\n\n"
            "## What changed\n\n"
            "1. Recommended title rewritten to lead with convergence:\n"
            "   - Old: \"" + OLD_DOCX_TITLE + "\"\n"
            "   - New: \"" + NEW_TITLE_EN + "\"\n"
            "2. Compressed Results / Discussion headlines rewritten to align "
            "with the four-step Cao et al. Nature 2026 PerturbFate "
            "narrative.\n"
            "3. Fig.5 rebuilt as \"K is the leverage point\" (4 panels): "
            "per-biome K ridge, beta invariance forest with 95% HDI band, "
            "HMP IBDMDB K-shift across control / UC / CD, beta invariance "
            "across 108 longitudinal IBD subjects.\n"
            "4. Graphical abstract v1 added: 16:9, 4-node narrative arc + 3 "
            "evidence rails (Primary EMP / Falsification / Replication) + 2 "
            "deliverable boxes (Theoretical: hierarchical universality; "
            "Translational: K-shift readout for disease, disturbance, "
            "habitat perturbation).\n\n"
            "## What did NOT change (intentional)\n\n"
            "- OSF preregistration v0.2 (`T5_OSF_preregistration_v0.2.md` and "
            "`.docx`) is unchanged.\n"
            "- All four OSF Step answer documents "
            "(`T5_OSF_Step{4,5,6,7}_*.docx`) are unchanged.\n"
            "- Preregistered hypotheses H1 through H7 are unchanged in "
            "wording, threshold, and pass / fail outcome.\n"
            "- v3 manuscript docx is preserved as a locked snapshot; the "
            "rewrite is in the new v3.1 file.\n\n"
            "## Why\n\n"
            "Cao et al. Nature 2026 (PerturbFate; "
            "doi:10.1038/s41586-026-10367-0) demonstrated that the four-step "
            "narrative \"different perturbations -> common state -> common "
            "program -> key nodes\" is a high-impact framing for convergent "
            "biology. The same evidentiary structure was already present in "
            "the T5 results (15 biomes -> shared Taylor backbone -> "
            "stochastic-logistic + Gamma AFD -> K-axis as leverage point), "
            "but had not been promoted to the title and graphical abstract. "
            "The 2026-05-04 repositioning is editorial, not analytical.\n\n"
            "## Provenance\n\n"
            "- Master sync script: `scripts/T5_sync_convergence_repositioning.py`\n"
            "- Fig.5 build script: `scripts/T5_fig5_leverage.py`\n"
            "- Graphical abstract build script: `scripts/T5_graphical_abstract.py`\n"
            "- v3 docx rewrite script: `scripts/T5_rewrite_title_abstract.py`\n"
        )
        p.write_text(body, encoding="utf-8")
        return p, ["  created post-hoc framing note"]
    return p, []


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main() -> None:
    steps = [
        ("manuscript v0.2 EN", step1_manuscript_en),
        ("manuscript v0.2 ZH", step2_manuscript_zh),
        ("cover letter v0.1", step3_cover_letter),
        ("v4 docx",            step4_v4_docx),
        ("ZH full-form docx",  step5_zh_full_docx),
        ("K承載力 ZH docx",    step6_zh_k_docx),
        ("figures/renamed copy", step7_copy_figures),
        ("FIGURE_DATA_PROVENANCE", step8_provenance_md),
        ("README.md",          step9_readme),
        ("wiki/log.md",        step10_log_md),
        ("wiki/t5_topic_frame.md", step11_topic_frame),
        ("wiki/submissions.md", step12_submissions_md),
        ("post-hoc framing note", step13_post_hoc_note),
    ]
    for name, fn in steps:
        try:
            target, log = fn()
            print(f"[OK] {name}: {target}")
            for line in log:
                print(line)
        except Exception as e:
            print(f"[ERR] {name}: {type(e).__name__}: {e}")


if __name__ == "__main__":
    main()
