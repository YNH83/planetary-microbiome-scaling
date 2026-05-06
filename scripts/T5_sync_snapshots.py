"""
Bring the 5 SNAPSHOT docx files into the 2026-05-04 evening sync standard.

Snapshots are intermediate versions intentionally retained for provenance,
but they should still carry the author block + canonical companion-document
pointers so any reader of any snapshot reaches the current submission state.

Targets (5 files):
    1. 04202026 A shared Taylor links biomes.docx
       (already has author block; needs companion-docs banner + archival note)
    2. T5_cover_letter_v0.2.docx
       (convergence title in body; missing author block)
    3. T5_title_intro_methods_results_discussion_package_v3.1_convergence.docx
       (title in body; missing author block; pre-figure-embedding snapshot)
    4. T5_title_intro_methods_results_discussion_package_v4_with_figures.docx
       (title in body; missing author block; pre-references-injection snapshot)
    5. T5_title_intro_methods_results_discussion_package_v4.1_with_references.docx
       (title in body; missing author block; pre-author-injection snapshot)

NOT touched (LOCKED per CLAUDE.md preregistration integrity rule):
    - T5_OSF_Step{4,5,6,7}_*.docx
    - T5_OSF_preregistration_v0.2.{md,docx}
    - T5_References_Verification_2026-04-29.docx (earlier baseline)
    - T5_title_intro_methods_results_discussion_package_v3.docx
      (preregistration-era snapshot, locked)

Idempotent: detects MARKER "[ARCHIVAL SNAPSHOT 2026-05-04 evening sync]" and
replaces in place.
"""
from __future__ import annotations
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/ynh83/Desktop/T5_Macroecology")

TARGETS = [
    (ROOT / "04202026 A shared Taylor links biomes.docx",
     "Author-block source fragment (2026-04-20). Carries the three-author block "
     "and three affiliations that were later injected into v4.2 manuscript and "
     "v0.3 cover letter."),
    (ROOT / "T5_cover_letter_v0.2.docx",
     "Cover letter snapshot before author-block injection (2026-05-04 morning). "
     "Replaced by T5_cover_letter_v0.3.docx in the evening sync."),
    (ROOT / "T5_title_intro_methods_results_discussion_package_v3.1_convergence.docx",
     "Manuscript snapshot at v3.1 (2026-05-04 morning), title rewritten to "
     "convergence framing but figures not yet embedded. Replaced by v4 (figures), "
     "v4.1 (references), v4.2 (authors)."),
    (ROOT / "T5_title_intro_methods_results_discussion_package_v4_with_figures.docx",
     "Manuscript snapshot at v4 (2026-05-04 morning), figures embedded but "
     "References section not yet appended. Replaced by v4.1 (references) and "
     "v4.2 (authors)."),
    (ROOT / "T5_title_intro_methods_results_discussion_package_v4.1_with_references.docx",
     "Manuscript snapshot at v4.1 (2026-05-04 evening), 40-entry References "
     "section appended but author block not yet injected. Replaced by v4.2 "
     "(authors)."),
]

MARKER = "[ARCHIVAL SNAPSHOT 2026-05-04 evening sync]"
CANONICAL_LATEST = "T5_title_intro_methods_results_discussion_package_v4.2_with_authors.docx"
CANONICAL_COVER = "T5_cover_letter_v0.3.docx"

CONVERGENCE_EN = (
    "Convergent Taylor scaling links planetary microbiomes through a "
    "habitat-modulated carrying-capacity axis"
)
AUTHOR_LINE = "Yu-Nan Huang [1,2,3], Pen-Hua Su [2,3,*], Chieh-Chen Huang [1,*]"
AFFILIATIONS = (
    "1. Department of Life Science, National Chung Hsing University, "
    "Taichung, Taiwan",
    "2. Division of Genetics and Metabolism, Department of Pediatrics, "
    "Chung Shan Medical University Hospital, Taichung, Taiwan",
    "3. School of Medicine, Chung Shan Medical University, Taichung, Taiwan",
)
CORRESPONDENCE = (
    "Pen-Hua Su, Tel: +886 4 2473 9595 (Ext. 21707), email: ninaphsu@gmail.com "
    "(ORCID 0000-0003-4174-5036)",
    "Chieh-Chen Huang, Tel: +886 4 2284 0416 (Ext. 405), email: "
    "cchuang@dragon.nchu.edu.tw (ORCID 0000-0002-3739-6315)",
)
ORCIDS = ("Yu-Nan Huang 0000-0003-1688-1685; Pen-Hua Su 0000-0003-4174-5036; "
          "Chieh-Chen Huang 0000-0002-3739-6315.")


def make_run(paragraph, text, *, bold=False, italic=False, size=10.5,
             color=None, east_asia=True):
    r = paragraph.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if east_asia:
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "PMingLiU")
    return r


def build_banner(stub, role_note):
    paras = []

    p = stub.add_paragraph()
    make_run(p, MARKER, bold=True, size=11, color="C00000")
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Latest manuscript: ", bold=True)
    make_run(p, CANONICAL_LATEST)
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Latest cover letter: ", bold=True)
    make_run(p, CANONICAL_COVER)
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Snapshot role: ", bold=True)
    make_run(p, role_note)
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Recommended title (English): ", bold=True)
    make_run(p, CONVERGENCE_EN)
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Authors: ", bold=True)
    make_run(p, AUTHOR_LINE)
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Affiliations:", bold=True)
    paras.append(p)
    for aff in AFFILIATIONS:
        p = stub.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        make_run(p, aff)
        paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Corresponding authors:", bold=True)
    paras.append(p)
    for c in CORRESPONDENCE:
        p = stub.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        make_run(p, c)
        paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "ORCIDs: ", bold=True)
    make_run(p, ORCIDS)
    paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "Companion documents:", bold=True)
    paras.append(p)
    for c in (
        "T5_References_FullField_Verification_2026-05-04.docx (40-entry "
        "verification across 12 lenses)",
        "SUBMISSION_CHECKLIST.md (10-section go / no-go at package root)",
        "drafts/T5_post_hoc_framing_note_2026-05-04.md (OSF preregistration "
        "integrity protection)",
    ):
        p = stub.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        make_run(p, "- " + c)
        paras.append(p)

    p = stub.add_paragraph()
    make_run(p, "OSF preregistration integrity: ", bold=True)
    make_run(p, "T5_OSF_preregistration_v0.2.{md,docx} and "
             "T5_OSF_Step{4,5,6,7}_*.docx are unchanged; H1 to H7 hypotheses "
             "and thresholds are unchanged. All 2026-05-04 changes are "
             "editorial repositioning, not analytical re-runs.")
    paras.append(p)

    p = stub.add_paragraph()  # trailing spacer
    paras.append(p)

    return paras


def remove_existing_banner(doc):
    paras = list(doc.paragraphs)
    start = None
    for i, p in enumerate(paras):
        if MARKER in p.text:
            start = i
            break
    if start is None:
        return False
    end = start
    seen_osf = False
    for j in range(start, min(start + 35, len(paras))):
        text = paras[j].text
        if text.startswith("OSF preregistration integrity"):
            seen_osf = True
            end = j
        elif seen_osf and text.strip() == "":
            end = j
            break
    for p in paras[start:end + 1]:
        p._element.getparent().remove(p._element)
    return True


def sync_one(target_path, role_note):
    doc = Document(str(target_path))
    removed = remove_existing_banner(doc)

    stub = Document()
    new_paras = build_banner(stub, role_note)

    # Insert after the first non-empty paragraph (the doc's own title)
    target_idx = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            target_idx = i
            break
    anchor_xml = doc.paragraphs[target_idx]._element
    for new_p in new_paras:
        anchor_xml.addnext(deepcopy(new_p._element))
        anchor_xml = anchor_xml.getnext()

    doc.save(str(target_path))
    return removed


if __name__ == "__main__":
    for target_path, role_note in TARGETS:
        before_size = target_path.stat().st_size
        replaced = sync_one(target_path, role_note)
        after_size = target_path.stat().st_size
        verb = "REPLACED" if replaced else "INSERTED"
        delta = after_size - before_size
        sign = "+" if delta >= 0 else ""
        print(f"{verb:8s}  {target_path.name}  "
              f"{before_size:,} -> {after_size:,} bytes  ({sign}{delta:,})")
    print("\nDone.")
