"""Convert T5_OSF_preregistration_v0.2.md to a formatted Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "drafts", "T5_OSF_preregistration_v0.2.docx")

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# --- Styles helper ---
def set_normal(para, size=11, bold=False, color=None):
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(14 if level == 1 else 12)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_para(doc, text, bold_prefix=None, indent=False):
    """Add a paragraph; bold_prefix highlights 'Key: ' style labels."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix and text.startswith(bold_prefix):
        label, _, rest = text.partition(" ")
        run = p.add_run(label + " ")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run2 = p.add_run(rest)
        run2.font.name = "Arial"
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return p

# ===== DOCUMENT BODY =====

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("T5 OSF Pre-registration v0.2")
run.font.name = "Arial"
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_paragraph()

# Study title
p = doc.add_paragraph()
r1 = p.add_run("Study title. ")
r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
r2 = p.add_run(
    "One macroecological law governs gut, soil, ocean, and air microbiomes: "
    "Bayesian hierarchical universality with invariant exponent and "
    "habitat-modulated carrying capacity."
)
r2.font.name = "Arial"; r2.font.size = Pt(11)

# Date
p = doc.add_paragraph()
r1 = p.add_run("Date of registration. ")
r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
r2 = p.add_run(
    "2026-04-17 (v0.1 locked 2026-04-16 morning primary thresholds H1 to H3; "
    "v0.2 adds secondary thresholds H4 to H7 before running the corresponding analyses)."
)
r2.font.name = "Arial"; r2.font.size = Pt(11)

doc.add_paragraph()

# --- Hypotheses ---
add_heading(doc, "Hypotheses (locked before data analysis)", level=1)

hypotheses = [
    ("H1.", "At least 8 of 15 EMPO-3 biomes will pass Taylor R\u00b2 \u2265 0.80 AND \u03b2 \u2208 [1.5, 2.5]."),
    ("H2.", "Universal vs biome-specific BIC difference \u2265 10 in favour of universal."),
    ("H3.", "Gamma abundance fluctuation distribution beats exponential in at least 70% of biomes."),
    ("H4 (v0.2).", "Bayesian hierarchical partial-pooling \u03b2_global posterior 95% HDI contains 2.0 AND PSIS-LOO rejects complete-pooling (\u0394ELPD > 3 SE)."),
    ("H5 (v0.2).", "At least 3 of 4 non-neutral null generators rejected at z > 5 (Hubbell, Fisher log-series, Preston lognormal, Shoemaker 2017 lognormal-neutral)."),
    ("H6 (v0.2).", "Shotgun-metagenomic replication satisfies qualitative Taylor bar (\u03b2 \u2208 [1.5, 2.5] AND R\u00b2 \u2265 0.80) in at least 2 of 3 curatedMG cohorts AND universal \u03b2 within 15% of EMP anchor of 1.966."),
    ("H7 (v0.2).", "Carrying-capacity K distribution differs across biomes (Kruskal-Wallis p < 0.01) AND per-biome \u03b2 coefficient of variation < 10%."),
]
for label, text in hypotheses:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(label + " ")
    r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = "Arial"; r2.font.size = Pt(11)

doc.add_paragraph()

# --- Primary analysis plan ---
add_heading(doc, "Primary Analysis Plan", level=1)

steps = [
    "Load EMP release 1 deblur 90 bp BIOM (raw data/emp/emp_deblur_90bp.release1.biom).",
    "Drop Hypersaline (n < 100 samples) and Plant corpus (insufficient taxa). 15 EMPO-3 biomes retained.",
    "Apply 20% prevalence filter per biome.",
    "Fit per-biome Taylor in log-log space by OLS, 1,000 residual-bootstrap CI.",
    "Pool data points, fit universal vs biome-specific OLS, compute BIC.",
    "Per-taxon Gamma vs exponential MLE + KS.",
    "Bayesian hierarchical PyMC NUTS with priors specified in Methods Section 4.5.",
    "Hubbell Etienne 2005 DM steady-state null, 3 x 3 grid on theta and migration, 90 replicates.",
    "Fisher log-series and Preston lognormal nulls via analytical steady state + resampling.",
    "Shoemaker 2017 lognormal-neutral with published parameters \u03bc_K = 2.5, \u03c3_K = 0.45, 90 replicates.",
    "curatedMetagenomicData HMP_2019_ibdmdb + NielsenHB_2014 + ZellerG_2014; same per-cohort Taylor + pooled BIC.",
    "Per-taxon K from (mean, var) pair under Grilli stochastic logistic; Kruskal-Wallis across biomes.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(step)
    run.font.name = "Arial"; run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run(
    "Pre-registered thresholds are primary outcomes. "
    "Any deviation from the pre-registration will be flagged in Results with rationale and marked as post-hoc."
)
run.font.name = "Arial"; run.font.size = Pt(11); run.italic = True

doc.add_paragraph()

# --- Sensitivity analyses ---
add_heading(doc, "Sensitivity Analyses (exploratory)", level=1)
p = doc.add_paragraph()
run = p.add_run(
    "Prevalence sweep {0.05 to 0.50}, rarefaction {1,000 to 20,000 reads}, "
    "taxonomy {ASV to phylum}, sample-size {500 to 26,181}, 150 bp read-length robustness. "
    "These are not part of the primary pre-registered decision."
)
run.font.name = "Arial"; run.font.size = Pt(11)

doc.add_paragraph()

# --- Anticipated failure modes ---
add_heading(doc, "Anticipated Failure Modes and Responses", level=1)

failures = [
    ("H1-H3 (decisive contrast):",
     "If any fails, demote the claim from universal to partial-universal in the manuscript."),
    ("H4 (hierarchical model):",
     "If hierarchical loses to universal-only, report complete-pooling as the final model and revise Discussion accordingly."),
    ("H5 (null falsification):",
     "If only Hubbell rejected, report that the neutral null is the only falsifiable alternative at reach of this data."),
    ("H6 (shotgun BIC):",
     "If shotgun BIC non-decisive, report honestly as a data-volume limitation and restrict the universality claim to 16S amplicon resolution."),
    ("H7 (K distribution):",
     "If fails, the 'host enters alpha not beta' Discussion claim is retracted to 'not yet derivable'."),
]
for label, text in failures:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(label + " ")
    r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = "Arial"; r2.font.size = Pt(11)

doc.add_paragraph()

# --- Data and code provenance ---
add_heading(doc, "Data and Code Provenance", level=1)

sources = [
    ("EMP release 1 deblur 90 bp BIOM:", "ftp.microbio.me/emp/release1/ (public)."),
    ("curatedMetagenomicData:", "Bioconductor package (public)."),
    ("Analysis scripts:", "scripts/T5_*.py under project repository (DOI pending)."),
]
for label, text in sources:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(label + " ")
    r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = "Arial"; r2.font.size = Pt(11)

doc.add_paragraph()

# --- Sign-off ---
add_heading(doc, "Sign-off", level=1)
p = doc.add_paragraph()
run = p.add_run(
    "To be finalised. All authors will review and sign this pre-registration "
    "before submission to Nature Ecology and Evolution."
)
run.font.name = "Arial"; run.font.size = Pt(11); run.italic = True

# Signature table
doc.add_paragraph()
table = doc.add_table(rows=6, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Author"
hdr[1].text = "Role"
hdr[2].text = "Date signed"
roles = [
    ("(First author)", "Conceptualization, Analysis, Writing (original draft)", ""),
    ("(Statistical advisor)", "Methodology, Formal analysis, Writing (review)", ""),
    ("(Microbiology expert / Corresponding)", "Conceptualization, Supervision, Writing (review)", ""),
    ("(International collaborator)", "Investigation, Writing (review)", ""),
    ("(Clinical physician)", "Clinical interpretation, Writing (review)", ""),
]
for i, (name, role, date) in enumerate(roles, 1):
    row = table.rows[i].cells
    row[0].text = name
    row[1].text = role
    row[2].text = date
for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)

doc.save(OUT)
print(f"Saved: {OUT}")
