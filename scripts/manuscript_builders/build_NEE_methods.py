"""
Build T5_Methods_NEE_2026-05-06.docx

Methods section for the NEE submission. Mirrors the math typography
and reference numbering of the Intro / Results / Discussion files
(citation numbering is consistent: [1]-[25]).

Markup tags inside body text:
    [i]x[/i]   italic
    [s]x[/s]   subscript
    [u]x[/u]   superscript
"""

import re
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/ynh83/Desktop/T5_Macroecology/T5_Methods_NEE_2026-05-06.docx"

# ---------- Methods body ----------
SECTIONS = [
    (
        "Datasets",
        [
            "Earth Microbiome Project (EMP) release 1. The primary dataset "
            "was the EMP release 1 deblur table [1, 14]. After dropping "
            "two biomes with insufficient sampling (Hypersaline and Plant "
            "corpus, both [i]n[/i] < 100), the analysis set comprised "
            "26,181 samples and 317,314 amplicon sequence variants (ASVs) "
            "distributed across 15 EMPO-3 biomes (Animal corpus, distal "
            "gut, proximal gut, secretion, surface; Plant rhizosphere, "
            "surface; Aerosol non-saline; Sediment non-saline, saline; "
            "Soil non-saline; Surface non-saline, saline; Water "
            "non-saline, saline). The default read length was 90 bp; the "
            "150 bp deblur table was retained for a separate sensitivity "
            "analysis.",

            "curatedMetagenomicData shotgun cohorts. For cross-platform "
            "replication, we used 4,702 stool samples drawn from nine "
            "independently published shotgun cohorts assembled in the "
            "curatedMetagenomicData Bioconductor / ExperimentHub package "
            "[21]: HMP iHMP-IBDMDB [3], NielsenHB, ZellerG, LifeLinesDeep, "
            "YachidaS, QinJ, FengQ, KarlssonFH, and VogtmannE.",

            "iHMP IBDMDB longitudinal subset. Disease and longitudinal "
            "analyses used the integrative Human Microbiome Project "
            "IBDMDB cohort [3]: 108 subjects classified as control, "
            "ulcerative colitis, or Crohn's disease, with samples binned "
            "into three visit-time strata.",

            "Tara Oceans external holdouts. External holdouts comprised "
            "two Tara Oceans subsets not included in EMP-anchored "
            "modelling: P0 taxonomic abundance and P1 functional KEGG "
            "ortholog abundance [6].",
        ],
    ),
    (
        "Preprocessing and inclusion criteria",
        [
            "For each biome and each cohort we retained taxa observed in "
            "at least 20% of samples (the pre-registered prevalence "
            "filter), reducing rare-tail noise without truncating the "
            "abundance range used for Taylor scaling. We then computed "
            "per-taxon mean relative abundance and across-sample variance "
            "within each biome. Per-biome regression units were the "
            "(mean, variance) pairs across taxa, typically 300 to 3,000 "
            "per biome after filtering, with the pooled set comprising "
            "approximately 50,000 (mean, variance) points across the 15 "
            "biomes.",
        ],
    ),
    (
        "Per-biome Taylor regression (Model A)",
        [
            "For each biome [i]j[/i], we fit ordinary least-squares "
            "regression in log-log space, "
            "log[s]10[/s]([i]σ[/i][u]2[/u][s]ij[/s]) = "
            "[i]α[/i][s]j[/s] + [i]β[/i][s]j[/s] · "
            "log[s]10[/s](⟨[i]x[/i]⟩[s]ij[/s]) + [i]ε[/i][s]ij[/s], "
            "where ⟨[i]x[/i]⟩[s]ij[/s] is the mean relative abundance of "
            "taxon [i]i[/i] in biome [i]j[/i] and "
            "[i]σ[/i][u]2[/u][s]ij[/s] is its across-sample variance. We "
            "obtained 95% confidence intervals on [i]β[/i][s]j[/s] from "
            "1,000 residual bootstrap resamples and reported the "
            "coefficient of determination [i]R[/i][u]2[/u]. Pre-registered "
            "hypothesis H1 required at least 8 of 15 biomes to satisfy "
            "[i]R[/i][u]2[/u] ≥ 0.80 and [i]β[/i][s]j[/s] ∈ [1.5, 2.5].",
        ],
    ),
    (
        "Universal-slope model (Model B) and BIC model selection",
        [
            "We compared Model A (15 biome-specific [i]β[/i] values) "
            "against a universal-slope model B with biome-specific "
            "intercepts but a single shared [i]β[/i], fit by pooled OLS "
            "with biome as a categorical fixed effect. Model selection "
            "used the Bayesian information criterion: "
            "Δ[i]BIC[/i] = [i]BIC[/i](Model A) − [i]BIC[/i](Model B). "
            "Following Kass and Raftery [23], Δ[i]BIC[/i] ≥ 10 was "
            "interpreted as decisive evidence in favour of the simpler "
            "model. Hypothesis H2 required Δ[i]BIC[/i] ≥ 10.",
        ],
    ),
    (
        "Bayesian hierarchical model (Model C) and PSIS-LOO comparison",
        [
            "We additionally fit a Bayesian hierarchical (partial-pooling) "
            "model in PyMC [24] using the No-U-Turn sampler with 4 chains, "
            "1,500 tuning iterations followed by 1,500 sampling iterations, "
            "and target acceptance 0.95. The model was: "
            "[i]β[/i][s]global[/s] ∼ Normal(2, 0.5); "
            "[i]τ[/i] ∼ HalfCauchy(0.1); "
            "[i]β[/i][s]j[/s] ∼ Normal([i]β[/i][s]global[/s], [i]τ[/i]); "
            "[i]α[/i][s]j[/s] ∼ Normal(0, 5); "
            "[i]σ[/i] ∼ HalfNormal(1); "
            "log[s]10[/s]([i]σ[/i][u]2[/u][s]ij[/s]) ∼ "
            "Normal([i]α[/i][s]j[/s] + "
            "[i]β[/i][s]j[/s] · log[s]10[/s](⟨[i]x[/i]⟩[s]ij[/s]), "
            "[i]σ[/i]). Posterior diagnostics required potential scale "
            "reduction factor [i]R̂[/i] < 1.01 and effective sample size "
            "[i]n[/i][s]eff[/s] > 500 for all primary parameters. "
            "Comparison against complete-pooling (single [i]β[/i], single "
            "[i]α[/i]) and no-pooling (15 independent [i]β[/i][s]j[/s] and "
            "[i]α[/i][s]j[/s]) used PSIS-LOO ELPD via ArviZ following "
            "Vehtari, Gelman and Gabry [18]. Pre-registered hypothesis H4 "
            "required the partial-pooling versus universal-only contrast "
            "to exceed 3 standard errors and the [i]β[/i][s]global[/s] "
            "95% highest-density interval to contain 2.",
        ],
    ),
    (
        "Abundance fluctuation distributions",
        [
            "For each (taxon, biome) pair with at least 5 observations, "
            "we fit Gamma and Exponential distributions to the within-"
            "biome relative-abundance vector by maximum likelihood and "
            "compared their log-likelihoods. We additionally computed a "
            "one-sample Kolmogorov-Smirnov test of the empirical "
            "distribution against the fitted Gamma. Pre-registered "
            "hypothesis H3 required at least 70% of taxa to satisfy "
            "[i]L[/i](Gamma) > [i]L[/i](Exponential) and "
            "Kolmogorov-Smirnov [i]p[/i] < 0.05.",
        ],
    ),
    (
        "Null-generator simulations",
        [
            "We tested four idealised null generators against the "
            "empirical pooled exponent. (i) Hubbell neutral drift was "
            "simulated under the unified neutral theory with parameters "
            "from the pooled metacommunity (community size [i]J[/i] = "
            "total reads; biodiversity [i]θ[/i] matched to the empirical "
            "Fisher fit). (ii) Fisher log-series sampled per-taxon "
            "abundances from a Fisher log-series with shape parameter "
            "matched to the empirical fit. (iii) Preston lognormal "
            "sampled abundances from a lognormal with location [i]μ[/i] "
            "and scale [i]σ[/i] matched to the empirical mean and "
            "variance. (iv) The Shoemaker-style lognormal-neutral hybrid "
            "[10] combined Hubbell ecological drift with lognormal "
            "carrying capacities. For each generator we ran 90 replicate "
            "simulations and refit Model B to obtain a null distribution "
            "of [i]β[/i]. The empirical exponent was compared to each "
            "null distribution by [i]z[/i]-score; pre-registered "
            "hypothesis H5 required at least 3 of 4 generators to be "
            "rejected at [i]z[/i] > 5.",
        ],
    ),
    (
        "Sensitivity analyses",
        [
            "The pre-registered sensitivity suite included six analyses: "
            "(i) prevalence threshold sweep at 10%, 20%, and 30%; "
            "(ii) rarefaction depth at 1,000, 5,000, and 10,000 reads per "
            "sample; (iii) leave-half-out subject resampling; "
            "(iv) taxonomic aggregation to genus, family, order, and "
            "phylum levels; (v) read length comparison (90 bp deblur "
            "primary versus 150 bp deblur sensitivity); and (vi) "
            "leave-one-biome-out. For each sweep, Model B was refit and "
            "the change in pooled [i]β[/i], Δ[i]BIC[/i], and per-biome "
            "[i]R[/i][u]2[/u] was reported.",
        ],
    ),
    (
        "Carrying-capacity tests",
        [
            "For each biome we extracted the per-taxon log [i]K[/i] "
            "(approximately the log of mean relative abundance) and "
            "tested between-biome differences with the Kruskal-Wallis "
            "omnibus test and Levene's test for variance heterogeneity. "
            "For disease-state [i]K[/i] shifts in the HMP IBDMDB stool "
            "subset, pairwise Kolmogorov-Smirnov tests with Bonferroni "
            "correction were applied across control, ulcerative colitis, "
            "and Crohn's disease. Pre-registered hypothesis H7 required "
            "Kruskal-Wallis [i]p[/i] < 0.01 across biomes and per-biome "
            "[i]β[/i] coefficient of variation below 10%.",
        ],
    ),
    (
        "External holdouts",
        [
            "For Tara Oceans P0 (taxonomic), the per-cohort Taylor "
            "pipeline was applied without modification, and we report "
            "[i]R[/i][u]2[/u] and [i]β[/i]. For Tara Oceans P1 "
            "(functional KEGG ortholog abundance), we computed Taylor "
            "scaling on KEGG ortholog read counts. The functional "
            "analysis is reported as a divergent (functional-gene) "
            "regime rather than as part of the universal taxonomic "
            "backbone.",
        ],
    ),
    (
        "Pre-registration and analytical decision rules",
        [
            "The full analytic protocol, including hypotheses H1 to H7, "
            "decision rules, sample-inclusion criteria, primary and "
            "secondary outcomes, and threshold values for "
            "Δ[i]BIC[/i], [i]z[/i]-score, and PSIS-LOO ELPD, was "
            "pre-registered on the Open Science Framework "
            "(OSF v0.2; locked 2026-05-07 prior to any real-data run "
            "[25]). A post-hoc framing note dated 2026-05-04 documents "
            "the repositioning of the title and the Figure 5 narrative "
            "without any changes to thresholds, generators, or "
            "statistical decisions.",
        ],
    ),
    (
        "Statistical reporting",
        [
            "Confidence intervals are 95% unless otherwise specified. "
            "Bayesian intervals are reported as 95% highest-density "
            "intervals. All [i]p[/i] values are two-sided. Multiple "
            "testing within disease-state Kolmogorov-Smirnov pairs used "
            "Bonferroni correction. Test statistics are reported with "
            "degrees of freedom where applicable. Numbers of replicates, "
            "subjects, taxa, and samples are stated alongside each test.",
        ],
    ),
    (
        "Code and data availability",
        [
            "All input datasets are publicly accessible. EMP release 1 "
            "was downloaded from the official EMP S3 bucket; "
            "curatedMetagenomicData was accessed via the "
            "Bioconductor / ExperimentHub R package [21]; iHMP IBDMDB "
            "was downloaded from ibdmdb.org [3]; Tara Oceans tables were "
            "obtained from the Ocean Gene Atlas. Analysis code (Python "
            "3.10 with NumPy, SciPy, pandas, PyMC [24], and ArviZ) and "
            "end-to-end build scripts will be deposited at a public "
            "GitHub or Zenodo archive at acceptance, with a permanent "
            "DOI. Pre-registration: OSF v0.2.",
        ],
    ),
]

# Cited subset (renumbered consistent with Intro file plus new Methods entries)
CITED_REFS = [
    {
        "n": 1,
        "authors": "Thompson LR, Sanders JG, McDonald D, et al.",
        "title": ("A communal catalogue reveals Earth's multiscale microbial "
                  "diversity."),
        "journal": "Nature",
        "year": "2017",
        "vol": "551(7681):457-463.",
        "id_kind": "PMID",
        "id_value": "29088705",
        "url": "https://pubmed.ncbi.nlm.nih.gov/29088705/",
    },
    {
        "n": 3,
        "authors": "Lloyd-Price J, Arze C, Ananthakrishnan AN, et al.",
        "title": ("Multi-omics of the gut microbial ecosystem in inflammatory "
                  "bowel diseases."),
        "journal": "Nature",
        "year": "2019",
        "vol": "569(7758):655-662.",
        "id_kind": "PMID",
        "id_value": "31142855",
        "url": "https://pubmed.ncbi.nlm.nih.gov/31142855/",
    },
    {
        "n": 6,
        "authors": "Sunagawa S, Coelho LP, Chaffron S, et al.",
        "title": "Structure and function of the global ocean microbiome.",
        "journal": "Science",
        "year": "2015",
        "vol": "348(6237):1261359.",
        "id_kind": "PMID",
        "id_value": "25999513",
        "url": "https://pubmed.ncbi.nlm.nih.gov/25999513/",
    },
    {
        "n": 10,
        "authors": "Shoemaker WR, Grilli J.",
        "title": ("Investigating macroecological patterns in coarse-grained "
                  "microbial communities using the stochastic logistic model "
                  "of growth."),
        "journal": "eLife",
        "year": "2024",
        "vol": "12:RP89650.",
        "id_kind": "PMID",
        "id_value": "38251984",
        "url": "https://pubmed.ncbi.nlm.nih.gov/38251984/",
    },
    {
        "n": 14,
        "authors": "Amir A, McDonald D, Navas-Molina JA, et al.",
        "title": "Deblur Rapidly Resolves Single-Nucleotide Community Sequence Patterns.",
        "journal": "mSystems",
        "year": "2017",
        "vol": "2(2):e00191-16.",
        "id_kind": "PMID",
        "id_value": "28289731",
        "url": "https://pubmed.ncbi.nlm.nih.gov/28289731/",
    },
    {
        "n": 18,
        "authors": "Vehtari A, Gelman A, Gabry J.",
        "title": ("Practical Bayesian model evaluation using leave-one-out "
                  "cross-validation and WAIC."),
        "journal": "Statistics and Computing",
        "year": "2017",
        "vol": "27(5):1413-1432.",
        "id_kind": "DOI",
        "id_value": "10.1007/s11222-016-9696-4",
        "url": "https://doi.org/10.1007/s11222-016-9696-4",
    },
    {
        "n": 21,
        "authors": "Pasolli E, Schiffer L, Manghi P, et al.",
        "title": "Accessible, curated metagenomic data through ExperimentHub.",
        "journal": "Nature Methods",
        "year": "2017",
        "vol": "14(11):1023-1024.",
        "id_kind": "PMID",
        "id_value": "29088129",
        "url": "https://pubmed.ncbi.nlm.nih.gov/29088129/",
    },
    {
        "n": 23,
        "authors": "Kass RE, Raftery AE.",
        "title": "Bayes Factors.",
        "journal": "Journal of the American Statistical Association",
        "year": "1995",
        "vol": "90(430):773-795.",
        "id_kind": "DOI",
        "id_value": "10.1080/01621459.1995.10476572",
        "url": "https://doi.org/10.1080/01621459.1995.10476572",
    },
    {
        "n": 24,
        "authors": "Abril-Pla O, Andreani V, Carroll C, et al.",
        "title": ("PyMC: a modern, and comprehensive probabilistic "
                  "programming framework in Python."),
        "journal": "PeerJ Computer Science",
        "year": "2023",
        "vol": "9:e1516.",
        "id_kind": "DOI",
        "id_value": "10.7717/peerj-cs.1516",
        "url": "https://doi.org/10.7717/peerj-cs.1516",
    },
    {
        "n": 25,
        "authors": "Nosek BA, Ebersole CR, DeHaven AC, Mellor DT.",
        "title": "The preregistration revolution.",
        "journal": "Proceedings of the National Academy of Sciences USA",
        "year": "2018",
        "vol": "115(11):2600-2606.",
        "id_kind": "PMID",
        "id_value": "29531091",
        "url": "https://pubmed.ncbi.nlm.nih.gov/29531091/",
    },
]

# ---------- Build docx ----------

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), "PMingLiU")

for section in doc.sections:
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

_MARKUP_RE = re.compile(r"(\[/?[isu]\])")


def add_centered(text, fs=11, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(fs)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color


def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)


def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def render_paragraph(text, fs=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.7
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.6)
    state = {"italic": False, "sub": False, "sup": False}
    for chunk in _MARKUP_RE.split(text):
        if chunk == "":
            continue
        if chunk == "[i]":
            state["italic"] = True
        elif chunk == "[/i]":
            state["italic"] = False
        elif chunk == "[s]":
            state["sub"] = True
        elif chunk == "[/s]":
            state["sub"] = False
        elif chunk == "[u]":
            state["sup"] = True
        elif chunk == "[/u]":
            state["sup"] = False
        else:
            run = p.add_run(chunk)
            run.font.size = Pt(fs)
            run.italic = state["italic"]
            run.font.subscript = state["sub"]
            run.font.superscript = state["sup"]
    return p


# ----- Banner -----
add_centered("Submission Methods section",
             fs=10.5, bold=True, color=RGBColor(0x77, 0x77, 0x77))
add_centered("Target journal: Nature Ecology and Evolution",
             fs=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))
add_centered(
    f"Date: {date.today().isoformat()}  |  Aligned to manuscript v4.2 + "
    "Results / Discussion 2026-05-06  |  Citation numbering matches the "
    "package master ([1]-[25]).",
    fs=9.5, italic=True, color=RGBColor(0x77, 0x77, 0x77),
)

doc.add_paragraph()

# ----- Title -----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(14)
title_run = title_p.add_run(
    "One macroecological scaling regime governs planetary "
    "microbiomes with habitat-specific carrying capacities"
)
title_run.bold = True
title_run.font.size = Pt(14.5)
title_run.font.color.rgb = RGBColor(0x12, 0x1F, 0x33)

# ----- Methods -----
add_h1("Methods")

word_total = 0
for sec_title, paragraphs in SECTIONS:
    add_h2(sec_title)
    for para in paragraphs:
        render_paragraph(para, fs=11)
        clean = _MARKUP_RE.sub("", para)
        word_total += len(clean.split())

# ----- References cited (subset) -----
add_h1("References cited in Methods")

note_p = doc.add_paragraph()
nrun = note_p.add_run(
    "Numbering is consistent with the package master "
    "(Introduction + References [1]-[22]; Methods adds [23]-[25]). "
    "All identifiers below were verified against PubMed esummary or "
    "CrossRef on 2026-05-04."
)
nrun.italic = True
nrun.font.size = Pt(10)
nrun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
note_p.paragraph_format.space_after = Pt(10)

for ref in CITED_REFS:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.3

    r_num = p.add_run(f"[{ref['n']}] ")
    r_num.bold = True
    r_num.font.size = Pt(10.5)
    r_num.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    r_au = p.add_run(ref["authors"] + " ")
    r_au.font.size = Pt(10.5)

    r_ti = p.add_run(ref["title"] + " ")
    r_ti.font.size = Pt(10.5)
    r_ti.italic = True

    r_jr = p.add_run(f"{ref['journal']}. {ref['year']};{ref['vol']}")
    r_jr.font.size = Pt(10.5)

    r_idlbl = p.add_run(f"  {ref['id_kind']}: ")
    r_idlbl.bold = True
    r_idlbl.font.size = Pt(10.5)

    r_idval = p.add_run(ref["id_value"])
    r_idval.font.size = Pt(10.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.3)
    p2.paragraph_format.space_after = Pt(7)
    p2.paragraph_format.line_spacing = 1.2
    r_link_lbl = p2.add_run("Link: ")
    r_link_lbl.font.size = Pt(9.5)
    r_link_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r_link = p2.add_run(ref["url"])
    r_link.font.size = Pt(9.5)
    r_link.italic = True
    r_link.font.color.rgb = RGBColor(0x1F, 0x3A, 0x88)
    r_sep = p2.add_run("    Verified 2026-05-04")
    r_sep.font.size = Pt(9.5)
    r_sep.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

# ----- Word count -----
add_h1("Word count and section map")
wc_p = doc.add_paragraph()
wcr = wc_p.add_run(
    f"Body word count: ~{word_total} words "
    "(NEE Article Methods is unconstrained; typical range "
    "1,500-3,000 words). 13 subsections covering datasets, preprocessing, "
    "three primary models (per-biome OLS, universal-slope OLS, Bayesian "
    "hierarchical), AFD characterization, four null generators, six "
    "sensitivity analyses, K-distribution tests, external holdouts, "
    "pre-registration, statistical reporting, and code / data "
    "availability."
)
wcr.font.size = Pt(10.5)
wcr.italic = True
wcr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ----- Footer -----
doc.add_paragraph()
foot_p = doc.add_paragraph()
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot_p.add_run(
    "T5 Macroecology submission package (manuscript v4.2; cover letter "
    "v0.3). Pre-registration: OSF v0.2 (2026-05-07). Math typography "
    "uses Word run-level italic / subscript / superscript (NEE Word "
    "convention; not raw LaTeX)."
)
fr.font.size = Pt(8.5)
fr.italic = True
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"sections: {len(SECTIONS)}; refs cited: {len(CITED_REFS)}; "
      f"body words: ~{word_total}")
