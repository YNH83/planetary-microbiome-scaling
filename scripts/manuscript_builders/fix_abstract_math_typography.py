"""
fix_abstract_math_typography.py

Re-render the Abstract paragraph in
'/Users/ynh83/Desktop/05062026 Taylor links biomes.docx' so that Greek
variables (β, α, K) and the R² superscript are formatted with proper
Word italic / superscript runs (matching Results / Discussion / Methods).

The text is preserved verbatim from the user's reorganised version
(US spelling 'idealized' / 'generalizable'); only run-level formatting
is added.
"""

import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

DOC_PATH = "/Users/ynh83/Desktop/05062026 Taylor links biomes.docx"

# Markup-tagged abstract: same text as user's saved version, with [i]/[u] tags.
ABSTRACT_MARKUP = (
    "Whether host-associated and free-living microbiomes obey the same "
    "ecological rules remains unresolved. Host-associated communities "
    "(gut, skin, oral cavity) are typically modelled through host "
    "filtering, immunity, and diet, whereas free-living microbiomes "
    "(soil, sediment, ocean, aerosol) are framed through environmental "
    "filtering, dispersal, and neutral drift. Using the Earth Microbiome "
    "Project release 1 deblur table (26,181 samples, 317,314 amplicon "
    "sequence variants, 15 EMPO-3 biomes), we tested whether a single "
    "macroecological scaling law links these domains. Taylor's law of "
    "variance-to-mean scaling held within every biome ([i]R[/i]² "
    "≥ 0.80), with a universal exponent [i]β[/i] = 1.950 (95% "
    "highest-density interval 1.909 to 1.992) preferred over "
    "biome-specific slopes (Δ[i]BIC[/i] = 25.7). 95% of taxa exhibited "
    "Gamma-shaped abundance fluctuation distributions consistent with "
    "stochastic logistic dynamics, and three of four idealized null "
    "generators (Hubbell neutral drift, Fisher log-series, Preston "
    "lognormal) were decisively falsified; a Shoemaker lognormal-neutral "
    "generator lay at the boundary. The backbone reproduced across nine "
    "shotgun metagenomic cohorts (4,702 samples; Δ[i]BIC[/i] = 23.4) and "
    "108 longitudinal inflammatory bowel disease subjects. Habitat, "
    "disease, and time perturbed only the carrying-capacity intercept "
    "[i]α[/i] (≈ log [i]K[/i]), leaving [i]β[/i] invariant. We propose "
    "[i]K[/i] as a generalizable readout linking ecological theory to "
    "microbiome perturbation, disturbance, and intervention."
)

_MARKUP_RE = re.compile(r"(\[/?[isu]\])")


def parse_markup(text):
    state = {"italic": False, "sub": False, "sup": False}
    for chunk in _MARKUP_RE.split(text):
        if not chunk:
            continue
        if chunk == "[i]":
            state["italic"] = True
        elif chunk == "[/i]":
            state["italic"] = False
        elif chunk == "[s]":
            state["sub"] = True
        elif chunk == "[/s]":
            state["sub"] = False
        elif chunk == "[u]":
            state["sup"] = True
        elif chunk == "[/u]":
            state["sup"] = False
        else:
            yield (chunk, state["italic"], state["sub"], state["sup"])


doc = Document(DOC_PATH)
abs_p = doc.paragraphs[8]

# Capture template rPr from existing first run (preserve font / size)
runs_xml = abs_p._element.findall(qn("w:r"))
template_rpr = None
for r in runs_xml:
    rpr = r.find(qn("w:rPr"))
    if rpr is not None:
        template_rpr = deepcopy(rpr)
        break

# Sanity check: existing text must match markup-stripped version (allow whitespace
# normalisation)
clean_markup = _MARKUP_RE.sub("", ABSTRACT_MARKUP)
existing = abs_p.text
if existing.strip() != clean_markup.strip():
    # Show diff
    import difflib
    print("WARN: existing abstract differs from markup-stripped version.")
    diff = difflib.unified_diff(
        existing.split(". "), clean_markup.split(". "),
        n=1, lineterm="")
    for line in list(diff)[:20]:
        print("  " + line)
    raise SystemExit("aborting; please align ABSTRACT_MARKUP to existing text.")

print("OK: existing abstract matches the markup version (stripped).")

# Clear existing runs but keep paragraph properties
for r in list(abs_p.runs):
    r._element.getparent().remove(r._element)

# Add new runs with italic / superscript / subscript
for text, italic, sub, sup in parse_markup(ABSTRACT_MARKUP):
    new_run = abs_p.add_run(text)
    if template_rpr is not None:
        existing_rpr = new_run._element.find(qn("w:rPr"))
        if existing_rpr is not None:
            new_run._element.remove(existing_rpr)
        new_run._element.insert(0, deepcopy(template_rpr))
    if italic:
        new_run.italic = True
    if sub:
        new_run.font.subscript = True
    if sup:
        new_run.font.superscript = True

doc.save(DOC_PATH)
print(f"saved: {DOC_PATH}")

# Verify
doc2 = Document(DOC_PATH)
p = doc2.paragraphs[8]
italic_count = sum(1 for r in p.runs if r.italic and r.text.strip())
sup_count = sum(1 for r in p.runs if r.font.superscript and r.text.strip())
print(f"after fix: italic runs = {italic_count}, superscript runs = {sup_count}")
print(f"abstract word count: {len(p.text.split())}")
