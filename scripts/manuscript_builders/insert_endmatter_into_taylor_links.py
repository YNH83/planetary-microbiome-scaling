"""
insert_endmatter_into_taylor_links.py

Insert the standard NEE end-matter sections into
/Users/ynh83/Desktop/05062026 Taylor links biomes.docx, following the
template observed in two recent NEE Articles
(s41559-026-03071-9 and s41559-026-03059-5):

  Before References: Reporting summary, Data availability, Code availability
  After References:  Acknowledgements, Author contributions, Funding,
                     Competing interests

Idempotent: skips any heading that already exists.
"""

import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC_PATH = "/Users/ynh83/Desktop/05062026 Taylor links biomes.docx"

PRE_REF_SECTIONS = [
    (
        "Reporting summary",
        [
            "Further information on research design is available in the "
            "Nature Portfolio Reporting Summary linked to this article."
        ],
    ),
    (
        "Data availability",
        [
            "All input datasets used in this study are publicly "
            "available. The Earth Microbiome Project release 1 deblur "
            "table is available from the official EMP archive. "
            "curatedMetagenomicData is available via the Bioconductor and "
            "ExperimentHub repositories. iHMP IBDMDB is available at "
            "https://ibdmdb.org. Tara Oceans tables are available from the "
            "Ocean Gene Atlas. Processed analytic outputs (per-biome "
            "moments, Bayesian posterior summaries, null-generator "
            "simulation outputs and figure source data) are deposited at "
            "the project GitHub repository "
            "(https://github.com/YNH83), with an archived Zenodo release "
            "(DOI assigned at acceptance). The pre-registration is "
            "publicly available on the Open Science Framework "
            "(T5 Macroecology v0.2; locked 2026-05-07)."
        ],
    ),
    (
        "Code availability",
        [
            "All analysis code (Python 3.10 with NumPy, SciPy, pandas, "
            "PyMC and ArviZ), build scripts and figure-generation "
            "pipelines are publicly available at "
            "https://github.com/YNH83 (T5_Macroecology repository), with "
            "an archived release at Zenodo upon publication."
        ],
    ),
]

POST_REF_SECTIONS = [
    (
        "Acknowledgements",
        [
            "[To be completed prior to submission.]"
        ],
    ),
    (
        "Author contributions",
        [
            "Y.N.H. led the conceptualization, methodology, data "
            "curation, formal analysis, software development, validation, "
            "visualization, and writing of the original draft. P.H.S. "
            "and C.C.H., as corresponding authors, contributed to "
            "conceptualization, supervised the project, obtained funding, "
            "and contributed to manuscript review and editing. All "
            "authors approved the submitted version and agree to be "
            "accountable for their own contributions and for the accuracy "
            "and integrity of the work."
        ],
    ),
    (
        "Funding",
        [
            "Research support was provided through the National Science "
            "and Technology Council, Taiwan (NSTC 113-2314-B-040-026-MY2 "
            "and NSTC 114-2622-B-040-001) and institutional grants from "
            "Chung Shan Medical University Hospital (CSH-2026-C-030)."
        ],
    ),
    (
        "Competing interests",
        [
            "The authors declare no competing interests."
        ],
    ),
]


# ---- Helpers (same as previous insert scripts) ----

_MARKUP_RE = re.compile(r"(\[/?[isu]\])")


def parse_markup(text):
    state = {"italic": False, "sub": False, "sup": False}
    for chunk in _MARKUP_RE.split(text):
        if not chunk:
            continue
        if chunk == "[i]":
            state["italic"] = True
        elif chunk == "[/i]":
            state["italic"] = False
        elif chunk == "[s]":
            state["sub"] = True
        elif chunk == "[/s]":
            state["sub"] = False
        elif chunk == "[u]":
            state["sup"] = True
        elif chunk == "[/u]":
            state["sup"] = False
        else:
            yield (chunk, state["italic"], state["sub"], state["sup"])


def _clone_with_no_runs(template_para):
    new_elem = deepcopy(template_para._element)
    for r in new_elem.findall(qn("w:r")):
        new_elem.remove(r)
    for h in new_elem.findall(qn("w:hyperlink")):
        new_elem.remove(h)
    return new_elem


def _capture_template_rpr(template_para):
    runs = template_para._element.findall(qn("w:r"))
    if not runs:
        return None
    rpr = runs[0].find(qn("w:rPr"))
    return deepcopy(rpr) if rpr is not None else None


def _add_run_with_template(new_para, text, template_rpr,
                           italic=False, sub=False, sup=False):
    r = new_para.add_run(text)
    if template_rpr is not None:
        existing = r._element.find(qn("w:rPr"))
        if existing is not None:
            r._element.remove(existing)
        r._element.insert(0, deepcopy(template_rpr))
    if italic:
        r.italic = True
    if sub:
        r.font.subscript = True
    if sup:
        r.font.superscript = True
    return r


def insert_paragraph_before(target_para, template_para, runs_data):
    new_elem = _clone_with_no_runs(template_para)
    template_rpr = _capture_template_rpr(template_para)
    target_para._element.addprevious(new_elem)
    new_para = Paragraph(new_elem, target_para._parent)
    for text, italic, sub, sup in runs_data:
        _add_run_with_template(new_para, text, template_rpr,
                                italic=italic, sub=sub, sup=sup)
    return new_para


def insert_paragraph_after(anchor_para, template_para, runs_data):
    new_elem = _clone_with_no_runs(template_para)
    template_rpr = _capture_template_rpr(template_para)
    anchor_para._element.addnext(new_elem)
    new_para = Paragraph(new_elem, anchor_para._parent)
    for text, italic, sub, sup in runs_data:
        _add_run_with_template(new_para, text, template_rpr,
                                italic=italic, sub=sub, sup=sup)
    return new_para


# ---- Open doc and find templates / anchors ----

doc = Document(DOC_PATH)
paras = doc.paragraphs

# Templates
template_h2 = None      # subsection heading style
template_body = None    # body paragraph style
ref_para = None         # 'Reference' h1
last_ref_entry = None   # last EndNote Bibliography entry

for p in paras:
    txt = p.text.strip()
    if template_h2 is None and txt == "Taylor's law holds within every EMPO-3 biome":
        template_h2 = p
    elif (template_body is None and
            txt.startswith("We first tested whether per-biome variance")):
        template_body = p
    elif ref_para is None and txt == "Reference":
        ref_para = p

# Find last bibliography entry
for p in paras:
    if p.style.name == "EndNote Bibliography" and p.text.strip():
        last_ref_entry = p

# Detect already-inserted sections (idempotency)
existing_headings = {p.text.strip() for p in paras if p.text.strip()}

print(f"templates: h2={template_h2 is not None} body={template_body is not None} "
      f"ref={ref_para is not None} last_ref={last_ref_entry is not None}")

# ---- Step 1: Insert PRE_REF_SECTIONS before 'Reference' ----

pre_inserted = 0
if ref_para and template_h2 and template_body:
    for sec_title, paragraphs in PRE_REF_SECTIONS:
        if sec_title in existing_headings:
            print(f"  pre  - SKIP (already present): {sec_title!r}")
            continue
        insert_paragraph_before(ref_para, template_h2, list(parse_markup(sec_title)))
        pre_inserted += 1
        for body_text in paragraphs:
            insert_paragraph_before(ref_para, template_body, list(parse_markup(body_text)))
            pre_inserted += 1
        # blank separator
        insert_paragraph_before(ref_para, template_body, [])
        pre_inserted += 1
        print(f"  pre  + inserted: {sec_title!r}")

# ---- Step 2: Insert POST_REF_SECTIONS after last bibliography entry ----

post_inserted = 0
if last_ref_entry and template_h2 and template_body:
    cursor = last_ref_entry
    for sec_title, paragraphs in POST_REF_SECTIONS:
        if sec_title in existing_headings:
            print(f"  post - SKIP (already present): {sec_title!r}")
            continue
        # Insert blank separator first (so heading sits visually below
        # the previous block)
        sep = insert_paragraph_after(cursor, template_body, [])
        cursor = sep
        post_inserted += 1
        # Insert heading
        h = insert_paragraph_after(cursor, template_h2,
                                    list(parse_markup(sec_title)))
        cursor = h
        post_inserted += 1
        # Insert body paragraphs
        for body_text in paragraphs:
            b = insert_paragraph_after(cursor, template_body,
                                        list(parse_markup(body_text)))
            cursor = b
            post_inserted += 1
        print(f"  post + inserted: {sec_title!r}")

# ---- Save ----

doc.save(DOC_PATH)
print(f"\nsaved: {DOC_PATH}")
print(f"summary: pre_inserted={pre_inserted}  post_inserted={post_inserted}")
