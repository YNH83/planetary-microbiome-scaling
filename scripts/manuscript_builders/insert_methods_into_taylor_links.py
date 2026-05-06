"""
insert_methods_into_taylor_links.py

Three operations on /Users/ynh83/Desktop/05062026 Taylor links biomes.docx:

1. Insert the Methods section (13 subsections, full math markup) just
   before the 'Reference' heading.
2. Repair the math typography of the existing Discussion body
   (replace plain Greek letters with italic Word runs by re-rendering
   each paragraph from build_NEE_discussion.py SECTIONS).
3. Append bibliography entries 22-25 in the existing EndNote-style
   plain-text format so Methods citations resolve.

Idempotent: skips Methods insertion if a 'Methods' h1 already exists,
and only adds bibliography entries that are missing.
"""

import re
import shutil
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC_PATH = "/Users/ynh83/Desktop/05062026 Taylor links biomes.docx"
METHODS_BUILD = "/Users/ynh83/Desktop/T5_Macroecology/build_NEE_methods.py"
DISCUSSION_BUILD = "/Users/ynh83/Desktop/T5_Macroecology/build_NEE_discussion.py"

# ---- Load SECTIONS from each build script (exec only the data prefix) ----

def load_sections(build_script):
    with open(build_script) as f:
        src = f.read()
    cutoff = src.find("# ---------- Build docx ----------")
    if cutoff < 0:
        cutoff = src.find("doc = Document()")
    ns = {}
    exec(src[:cutoff], ns)
    return ns

ns_m = load_sections(METHODS_BUILD)
METHODS_SECTIONS = ns_m["SECTIONS"]
METHODS_REFS = ns_m["CITED_REFS"]

ns_d = load_sections(DISCUSSION_BUILD)
DISCUSSION_SECTIONS = ns_d["SECTIONS"]

print(f"loaded {len(METHODS_SECTIONS)} Methods sections, "
      f"{len(METHODS_REFS)} Methods refs, "
      f"{len(DISCUSSION_SECTIONS)} Discussion sections")

_MARKUP_RE = re.compile(r"(\[/?[isu]\])")


def parse_markup(text):
    state = {"italic": False, "sub": False, "sup": False}
    for chunk in _MARKUP_RE.split(text):
        if not chunk:
            continue
        if chunk == "[i]":
            state["italic"] = True
        elif chunk == "[/i]":
            state["italic"] = False
        elif chunk == "[s]":
            state["sub"] = True
        elif chunk == "[/s]":
            state["sub"] = False
        elif chunk == "[u]":
            state["sup"] = True
        elif chunk == "[/u]":
            state["sup"] = False
        else:
            yield (chunk, state["italic"], state["sub"], state["sup"])


# ---- Insertion / replacement helpers preserving styling from a template ----

def _clone_with_no_runs(template_para):
    """Return a fresh <w:p> XML element with same paragraph properties as
    template_para but with all runs removed (so we can re-add them)."""
    new_elem = deepcopy(template_para._element)
    for r in new_elem.findall(qn("w:r")):
        new_elem.remove(r)
    # also clear hyperlink etc.
    for h in new_elem.findall(qn("w:hyperlink")):
        new_elem.remove(h)
    return new_elem


def _capture_template_rpr(template_para):
    """Capture run properties (rPr) from the first run of the template
    paragraph, so we can replicate font/size/bold/etc. on new runs."""
    runs = template_para._element.findall(qn("w:r"))
    if not runs:
        return None
    rpr = runs[0].find(qn("w:rPr"))
    return deepcopy(rpr) if rpr is not None else None


def _add_run_with_template(new_para, text, template_rpr,
                           italic=False, sub=False, sup=False):
    r = new_para.add_run(text)
    if template_rpr is not None:
        existing = r._element.find(qn("w:rPr"))
        if existing is not None:
            r._element.remove(existing)
        r._element.insert(0, deepcopy(template_rpr))
    if italic:
        r.italic = True
    if sub:
        r.font.subscript = True
    if sup:
        r.font.superscript = True
    return r


def insert_paragraph_before(target_para, template_para, runs_data):
    """Insert a new paragraph before target_para, replicating
    template_para's pPr and rPr. Adds runs with [i]/[s]/[u] formatting."""
    new_elem = _clone_with_no_runs(template_para)
    template_rpr = _capture_template_rpr(template_para)
    target_para._element.addprevious(new_elem)
    new_para = Paragraph(new_elem, target_para._parent)
    for text, italic, sub, sup in runs_data:
        _add_run_with_template(new_para, text, template_rpr,
                                italic=italic, sub=sub, sup=sup)
    return new_para


def replace_paragraph_runs(paragraph, runs_data):
    """Clear runs of paragraph and rebuild from runs_data, preserving
    paragraph-level pPr and the font properties of the first existing run."""
    template_rpr = _capture_template_rpr(paragraph)
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    for text, italic, sub, sup in runs_data:
        _add_run_with_template(paragraph, text, template_rpr,
                                italic=italic, sub=sub, sup=sup)


# ---- Begin operation ----

doc = Document(DOC_PATH)
paras = doc.paragraphs

# Detect existing 'Methods' heading (skip insertion if already done)
methods_already = any(p.text.strip() == "Methods" for p in paras)
if methods_already:
    print("Methods heading already present; skipping insertion step.")

# Find templates
template_h1 = None
template_h2 = None
template_body = None
ref_para = None

for p in paras:
    txt = p.text.strip()
    if template_h1 is None and txt == "Results":
        template_h1 = p
    elif template_h2 is None and txt == "Taylor's law holds within every EMPO-3 biome":
        template_h2 = p
    elif (template_body is None and
            txt.startswith("We first tested whether per-biome variance")):
        template_body = p
    elif ref_para is None and txt == "Reference":
        ref_para = p

print(f"templates: h1={template_h1 is not None}  h2={template_h2 is not None}  "
      f"body={template_body is not None}  ref={ref_para is not None}")

# ---- Step 1: Insert Methods section ----

methods_inserted = 0
if not methods_already and ref_para and template_h1 and template_h2 and template_body:
    # 'Methods' h1
    insert_paragraph_before(ref_para, template_h1,
                             list(parse_markup("Methods")))
    methods_inserted += 1

    for sec_title, paragraphs in METHODS_SECTIONS:
        insert_paragraph_before(ref_para, template_h2,
                                 list(parse_markup(sec_title)))
        methods_inserted += 1
        for body_text in paragraphs:
            insert_paragraph_before(ref_para, template_body,
                                     list(parse_markup(body_text)))
            methods_inserted += 1

    # blank separator
    blank_runs = []  # empty
    insert_paragraph_before(ref_para, template_body, blank_runs)
    methods_inserted += 1
    print(f"step 1: inserted {methods_inserted} Methods paragraphs")
else:
    print("step 1: skipped")

# Refresh paragraph list since we inserted nodes
doc_paras = doc.paragraphs

# ---- Step 2: Repair Discussion math typography ----

# Map existing Discussion bodies to my SECTIONS by sequence. Discussion
# in target file is paras after 'Discussion' heading and before next h1.
disc_h1_idx = None
for i, p in enumerate(doc_paras):
    if p.text.strip() == "Discussion":
        disc_h1_idx = i
        break

discussion_replaced = 0
if disc_h1_idx is not None:
    # Collect Discussion body paragraphs (until next non-empty paragraph that
    # is a known h1 like 'Methods', 'Reference', or until end).
    disc_bodies = []
    j = disc_h1_idx + 1
    while j < len(doc_paras):
        txt = doc_paras[j].text.strip()
        if txt in ("Methods", "Reference", ""):
            j += 1
            if txt in ("Methods", "Reference"):
                break
            continue
        disc_bodies.append(doc_paras[j])
        j += 1

    # Flatten my SECTIONS bodies into a list
    my_disc_bodies = []
    for sec_title, paragraphs in DISCUSSION_SECTIONS:
        for body in paragraphs:
            my_disc_bodies.append(body)

    print(f"step 2: existing disc body count = {len(disc_bodies)}, "
          f"my disc body count = {len(my_disc_bodies)}")

    if len(disc_bodies) == len(my_disc_bodies):
        for existing_p, my_text in zip(disc_bodies, my_disc_bodies):
            replace_paragraph_runs(existing_p, list(parse_markup(my_text)))
            discussion_replaced += 1
        print(f"step 2: replaced {discussion_replaced} Discussion bodies "
              "with markup-formatted versions")
    else:
        print("step 2: WARN body counts differ; skipping Discussion repair")
else:
    print("step 2: no Discussion section found; skipping")

# ---- Step 3: Append refs 22-25 if missing ----

# Find last EndNote Bibliography paragraph
last_ref_p = None
existing_ref_nums = set()
for p in doc.paragraphs:
    if p.style.name == "EndNote Bibliography":
        last_ref_p = p
        m = re.match(r"^\s*(\d+)\b", p.text)
        if m:
            existing_ref_nums.add(int(m.group(1)))

print(f"step 3: existing ref numbers up to {max(existing_ref_nums) if existing_ref_nums else 0}; "
      f"set has {len(existing_ref_nums)} entries")

NEW_REFS = {
    22: ("Verhulst, P. F. Notice sur la loi que la population poursuit "
         "dans son accroissement. Correspondance Mathematique et Physique "
         "10, 113-121 (1838)."),
    23: ("Kass, R. E. & Raftery, A. E. Bayes Factors. Journal of the "
         "American Statistical Association 90, 773-795 (1995). "
         "https://doi.org/10.1080/01621459.1995.10476572"),
    24: ("Abril-Pla, O. et al. PyMC: a modern, and comprehensive "
         "probabilistic programming framework in Python. PeerJ Comput Sci "
         "9, e1516 (2023). https://doi.org/10.7717/peerj-cs.1516"),
    25: ("Nosek, B. A., Ebersole, C. R., DeHaven, A. C. & Mellor, D. T. "
         "The preregistration revolution. Proc Natl Acad Sci U S A 115, "
         "2600-2606 (2018). https://doi.org/10.1073/pnas.1708274114"),
}

refs_added = 0
to_add = sorted(set(NEW_REFS) - existing_ref_nums)
if last_ref_p is not None and to_add:
    template_rpr = _capture_template_rpr(last_ref_p)
    cursor = last_ref_p
    for num in to_add:
        new_elem = _clone_with_no_runs(cursor)
        cursor._element.addnext(new_elem)
        new_para = Paragraph(new_elem, cursor._parent)
        text = f"{num}\t{NEW_REFS[num]}"
        _add_run_with_template(new_para, text, template_rpr)
        cursor = new_para
        refs_added += 1
    print(f"step 3: appended {refs_added} bibliography entries: {to_add}")
else:
    print("step 3: nothing to append")

# ---- Save ----

doc.save(DOC_PATH)
print(f"\nsaved: {DOC_PATH}")
print(f"summary: methods_inserted={methods_inserted}  "
      f"discussion_replaced={discussion_replaced}  refs_added={refs_added}")
