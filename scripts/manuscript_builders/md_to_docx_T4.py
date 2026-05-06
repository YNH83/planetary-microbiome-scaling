"""
Convert T4_ZH_first_draft.md to Word document on Desktop.

Handles: headings (#..####), blockquotes (>), tables (|...|),
bullets (-), horizontal rules (---), inline bold (**), italic (*),
and inline code (`).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(
    "/Users/ynh83/Desktop/Genetics 新題目跟架構 in UKB and All of us/"
    "04182026 新研究主題規劃/manuscripts/T4_ZH_first_draft.md"
)
OUT = Path("/Users/ynh83/Desktop/T4_ZH_first_draft.docx")

CJK_FONT = "PingFang TC"
LATIN_FONT = "Arial"
MONO_FONT = "Menlo"

HEADING_COLORS = {1: "3C5488", 2: "E64B35", 3: "00A087", 4: "3C5488"}
HEADING_SIZES = {1: 19, 2: 15, 3: 13, 4: 12}


# ---- low level helpers ----

def set_run_font(run, size=11, bold=False, italic=False, color=None, mono=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    ascii_font = MONO_FONT if mono else LATIN_FONT
    run.font.name = ascii_font
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def shade_cell(cell, hexcolor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tc_pr.append(shd)


def shade_paragraph(p, hexcolor):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    ppr.append(shd)


# ---- inline parser ----

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>[^*\n]+?)\*\*)"      # **bold**
    r"|(\*(?P<italic>[^*\n]+?)\*)"        # *italic*
    r"|(`(?P<code>[^`\n]+?)`)"            # `code`
)


def render_inline(paragraph, text, base_size=11, base_bold=False, base_italic=False,
                  base_color=None):
    """Parse markdown inline markup and emit runs on paragraph."""
    idx = 0
    for m in INLINE_RE.finditer(text):
        start, end = m.span()
        if start > idx:
            r = paragraph.add_run(text[idx:start])
            set_run_font(r, size=base_size, bold=base_bold, italic=base_italic,
                         color=base_color)
        if m.group("bold") is not None:
            r = paragraph.add_run(m.group("bold"))
            set_run_font(r, size=base_size, bold=True, italic=base_italic,
                         color=base_color)
        elif m.group("italic") is not None:
            r = paragraph.add_run(m.group("italic"))
            set_run_font(r, size=base_size, bold=base_bold, italic=True,
                         color=base_color)
        elif m.group("code") is not None:
            r = paragraph.add_run(m.group("code"))
            set_run_font(r, size=base_size - 1, mono=True, color="3C5488")
        idx = end
    if idx < len(text):
        r = paragraph.add_run(text[idx:])
        set_run_font(r, size=base_size, bold=base_bold, italic=base_italic,
                     color=base_color)


# ---- block emitters ----

def add_heading(doc, text, level):
    level = min(level, 4)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    render_inline(p, text,
                  base_size=HEADING_SIZES[level],
                  base_bold=True,
                  base_color=HEADING_COLORS[level])


def add_paragraph(doc, text, space_after=4, line_spacing=1.45):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    render_inline(p, text)


def add_blockquote(doc, lines):
    """Render consecutive > lines as a single shaded block."""
    text = "\n".join(lines).strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    shade_paragraph(p, "FFF8E7")
    render_inline(p, text, base_size=10, base_color="7E6148")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.35
    render_inline(p, text)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(doc, rows):
    """rows: list of list-of-cell-strings, first row = header."""
    ncol = max(len(r) for r in rows)
    padded = [r + [""] * (ncol - len(r)) for r in rows]
    t = doc.add_table(rows=len(padded), cols=ncol)
    t.style = "Light Grid Accent 1"
    for ci, cell_text in enumerate(padded[0]):
        c = t.rows[0].cells[ci]
        c.text = ""
        shade_cell(c, "3C5488")
        p = c.paragraphs[0]
        render_inline(p, cell_text, base_size=10, base_bold=True, base_color="FFFFFF")
    for ri in range(1, len(padded)):
        for ci, cell_text in enumerate(padded[ri]):
            c = t.rows[ri].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            render_inline(p, cell_text, base_size=10)


# ---- parser: line-by-line state machine ----

def convert(md_text: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    # Default style: CJK via PingFang TC, Latin via Arial
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)

    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    blockquote_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_blockquote():
        nonlocal blockquote_buf
        if blockquote_buf:
            add_blockquote(doc, blockquote_buf)
            blockquote_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            # drop markdown separator row |---|---|
            cleaned = [r for r in table_buf
                       if not all(re.fullmatch(r":?-+:?", c.strip()) for c in r)]
            add_table(doc, cleaned)
            table_buf = []

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # blockquote
        if line.startswith(">"):
            flush_table()
            content = line[1:].lstrip()
            blockquote_buf.append(content)
            i += 1
            continue
        else:
            flush_blockquote()

        # table
        if "|" in line and line.lstrip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # empty
        if not line.strip():
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", line.strip()):
            add_hr(doc)
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # bullets
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            add_bullet(doc, text)
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.35
            render_inline(p, text)
            i += 1
            continue

        # normal paragraph: accumulate consecutive non-empty non-special lines
        para_lines = [line]
        j = i + 1
        while j < n:
            nxt = lines[j].rstrip()
            if (not nxt.strip()
                or nxt.startswith("#")
                or nxt.startswith(">")
                or nxt.lstrip().startswith("|")
                or re.fullmatch(r"-{3,}", nxt.strip())
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)):
                break
            para_lines.append(nxt)
            j += 1
        add_paragraph(doc, " ".join(para_lines))
        i = j

    flush_blockquote()
    flush_table()

    return doc


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    md_text = SRC.read_text(encoding="utf-8")
    doc = convert(md_text)
    doc.save(OUT)
    print(f"[done] wrote {OUT}")
    print(f"[size] {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
