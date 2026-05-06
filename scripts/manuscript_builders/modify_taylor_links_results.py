"""
modify_taylor_links_results.py

Open `/Users/ynh83/Desktop/05062026 Taylor links biomes.docx` and replace
the Results body paragraphs with run-level italic / subscript / superscript
formatting (NEE math typography). Only the Results section is touched;
title, abstract, introduction, and references are left as-is.

SECTIONS data is reused from build_NEE_results.py at runtime (no duplication).
"""

import re
from docx import Document

DOC_PATH = "/Users/ynh83/Desktop/05062026 Taylor links biomes.docx"
BUILD_SCRIPT = "/Users/ynh83/Desktop/T5_Macroecology/build_NEE_results.py"

# ---- Load SECTIONS from build script (exec only the prefix that has data) ----
with open(BUILD_SCRIPT) as f:
    src = f.read()
cutoff = src.find("# ---------- Build docx ----------")
if cutoff < 0:
    cutoff = src.find("doc = Document()")
ns = {}
exec(src[:cutoff], ns)
SECTIONS = ns["SECTIONS"]
print(f"loaded {len(SECTIONS)} sections from build script")

# ---- Markup parser (NEE Word math typography) ----
_MARKUP_RE = re.compile(r"(\[/?[isu]\])")


def parse_markup(text):
    """Yield (text, italic, sub, sup) tuples honouring [i] [s] [u] tags."""
    state = {"italic": False, "sub": False, "sup": False}
    for chunk in _MARKUP_RE.split(text):
        if not chunk:
            continue
        if chunk == "[i]":
            state["italic"] = True
        elif chunk == "[/i]":
            state["italic"] = False
        elif chunk == "[s]":
            state["sub"] = True
        elif chunk == "[/s]":
            state["sub"] = False
        elif chunk == "[u]":
            state["sup"] = True
        elif chunk == "[/u]":
            state["sup"] = False
        else:
            yield (chunk, state["italic"], state["sub"], state["sup"])


def replace_paragraph_runs(paragraph, segments):
    """Clear paragraph runs and rebuild with formatting, preserving the
    paragraph-level properties (alignment, spacing, indent, style)."""
    # Capture original run font size if any (preserve sizing)
    fs = None
    for r in paragraph.runs:
        if r.font.size:
            fs = r.font.size
            break
    # Remove existing runs (but keep <w:pPr> and other paragraph-level XML)
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    # Add new runs with run-level formatting
    for text, italic, sub, sup in segments:
        run = paragraph.add_run(text)
        if fs is not None:
            run.font.size = fs
        if italic:
            run.italic = True
        if sub:
            run.font.subscript = True
        if sup:
            run.font.superscript = True


# ---- Open target docx and walk paragraphs ----
doc = Document(DOC_PATH)
section_map = {heading: bodies for heading, bodies in SECTIONS}
paras = doc.paragraphs

i = 0
replaced = 0
section_replacements = []
while i < len(paras):
    p = paras[i]
    txt = p.text.strip()
    if txt in section_map:
        bodies = section_map[txt]
        j = i + 1
        body_used = 0
        while body_used < len(bodies) and j < len(paras):
            # Skip empty paragraphs between heading and body
            if not paras[j].text.strip():
                j += 1
                continue
            body_text = bodies[body_used]
            segs = list(parse_markup(body_text))
            replace_paragraph_runs(paras[j], segs)
            replaced += 1
            section_replacements.append((txt[:60], j))
            body_used += 1
            j += 1
        i = j
    else:
        i += 1

print(f"\nreplaced {replaced} body paragraphs across {len(section_map)} sections:")
for sec, idx in section_replacements:
    print(f"  para[{idx:3d}] :: {sec}")

doc.save(DOC_PATH)
print(f"\nsaved: {DOC_PATH}")
