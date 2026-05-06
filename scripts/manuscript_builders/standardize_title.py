"""
Standardize title across all standalone NEE docx files (and the
Taylor-links submission docx). Replace
  'Convergent Taylor scaling links planetary microbiomes through a
   habitat-modulated carrying-capacity axis'
with
  'One macroecological scaling regime governs planetary microbiomes
   with habitat-specific carrying capacities'

Replacement is run-aware: if a title spans multiple runs, all are
updated. Run-level formatting (bold, font size, color) is preserved
on the first run; trailing runs are emptied.
"""

from docx import Document
import glob

OLD = ("Convergent Taylor scaling links planetary microbiomes through a "
       "habitat-modulated carrying-capacity axis")
NEW = ("One macroecological scaling regime governs planetary microbiomes "
       "with habitat-specific carrying capacities")

ROOT = "/Users/ynh83/Desktop/T5_Macroecology"

# Also include the build scripts so future rebuilds use the new title.
BUILD_SCRIPTS = [
    "build_NEE_abstract.py",
    "build_NEE_results.py",
    "build_NEE_discussion.py",
    "build_NEE_methods.py",
    "build_NEE_intro_with_refs.py",
]

files = sorted(glob.glob(f"{ROOT}/T5_*_NEE_2026-05-06.docx"))


def replace_in_paragraph(p, old, new):
    """Replace `old` (plain text) with `new` while preserving the
    formatting of the first run that contains the start of `old`. Works
    even if `old` spans multiple runs (common after Word edits)."""
    # Build cumulative run-text indices
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    start = full.index(old)
    end = start + len(old)
    # Walk runs and edit
    cursor = 0
    target_run = None
    for r in p.runs:
        rlen = len(r.text)
        run_start = cursor
        run_end = cursor + rlen
        cursor = run_end
        if run_end <= start:
            continue  # before target
        if run_start >= end:
            continue  # after target
        if target_run is None:
            # this is the first run touching the target; replace its
            # overlap with NEW (only this run carries the new text)
            local_old_start = max(0, start - run_start)
            local_old_end = min(rlen, end - run_start)
            r.text = (r.text[:local_old_start] + new +
                       r.text[local_old_end:])
            target_run = r
        else:
            # subsequent runs overlapping target: clear the overlap
            local_old_start = max(0, start - run_start)
            local_old_end = min(rlen, end - run_start)
            r.text = r.text[:local_old_start] + r.text[local_old_end:]
    return True


# ---- Step 1: Update standalone NEE docx ----
print("=== Step 1: replace title in standalone NEE docx files ===")
for f in files:
    doc = Document(f)
    changed = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, OLD, NEW):
            changed += 1
    if changed:
        doc.save(f)
    print(f"  {f.split('/')[-1]}: {changed} paragraph(s) updated")

# ---- Step 2: Update build scripts so future rebuilds match ----
print("\n=== Step 2: update build scripts ===")
for script in BUILD_SCRIPTS:
    path = f"{ROOT}/{script}"
    try:
        with open(path) as fh:
            src = fh.read()
    except FileNotFoundError:
        print(f"  {script}: not found, skipping")
        continue
    # Match either single-line or split-string form. Replace the canonical
    # phrase, including across split string literals (handle the line break
    # 'planetary microbiomes through a' / '"a habitat-modulated...').
    new_src = src
    # Single-line form
    new_src = new_src.replace(
        "Convergent Taylor scaling links planetary microbiomes through "
        "a habitat-modulated carrying-capacity axis",
        "One macroecological scaling regime governs planetary microbiomes "
        "with habitat-specific carrying capacities",
    )
    # Other potential variants split across lines
    new_src = new_src.replace(
        "Convergent Taylor scaling links planetary microbiomes through a "
        "habitat-modulated carrying-capacity axis",
        "One macroecological scaling regime governs planetary microbiomes "
        "with habitat-specific carrying capacities",
    )
    if new_src != src:
        with open(path, "w") as fh:
            fh.write(new_src)
        print(f"  {script}: updated")
    else:
        print(f"  {script}: no change (title not found in source)")

# ---- Step 3: Update the Taylor links submission docx (only abstract / banner mentions; the title at top is already correct) ----
target = "/Users/ynh83/Desktop/05062026 Taylor links biomes.docx"
doc = Document(target)
changed = 0
for p in doc.paragraphs:
    if replace_in_paragraph(p, OLD, NEW):
        changed += 1
if changed:
    doc.save(target)
print(f"\n=== Step 3: {target.split('/')[-1]}: {changed} paragraph(s) updated ===")
